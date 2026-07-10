from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Depends
from fastapi.responses import StreamingResponse, HTMLResponse
from pydantic import BaseModel
from typing import Optional
import uuid
import os
import io
import re
import json
from datetime import datetime
from urllib.parse import quote
from openai import APITimeoutError, APIConnectionError, APIStatusError, RateLimitError, AuthenticationError

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.api.endpoints.auth import get_current_user, get_current_user_profile
from app.core.school_data import load_classes, load_class_types
from app.core.user_data import get_user_subject, get_materials_path, get_library_trash_path, get_user_config_path
from app.core.paths import UPLOAD_DIR
from app.utils.file_lock import read_json, write_json
from app.utils.doc_converter import convert_docx_to_pdf, _convert_sync
from app.utils.ai_client import create_client, KIMI_BASE_URL
import logging

router = APIRouter()


MAGIC_JPEG = b"\xff\xd8\xff"
MAGIC_PNG = b"\x89\x50\x4e\x47"
MAGIC_PDF = b"%PDF"
MAGIC_ZIP = b"PK\x03\x04"
DOC_TYPES = {"题目", "答案", "词单", "资料"}


class MaterialTypeUpdateRequest(BaseModel):
    doc_type: str
    reason: str = ""


def _material_tree_node(name: str = "") -> dict:
    return {
        "name": name,
        "path": name,
        "count": 0,
        "children": {},
        "items": [],
    }


def _build_tag_tree(materials: list[dict]) -> list[dict]:
    root = _material_tree_node("")
    for material in materials:
        tag = (material.get("tag") or material.get("time") or "未分类").strip("/")
        parts = [part.strip() for part in tag.split("/") if part.strip()] or ["未分类"]
        cursor = root
        prefix = []
        for part in parts:
            prefix.append(part)
            if part not in cursor["children"]:
                cursor["children"][part] = {
                    "name": part,
                    "path": "/".join(prefix),
                    "count": 0,
                    "children": {},
                    "items": [],
                }
            cursor = cursor["children"][part]
            cursor["count"] += 1
        cursor["items"].append({
            "id": material["id"],
            "filename": material["filename"],
            "file_type": material.get("file_type", "pdf"),
            "created_at": material.get("created_at", ""),
            "file_size": material.get("file_size", 0),
        })

    def serialize(node: dict) -> dict:
        return {
            "name": node["name"],
            "path": node["path"],
            "count": node["count"],
            "items": node["items"],
            "children": [serialize(child) for child in sorted(node["children"].values(), key=lambda item: item["name"])],
        }

    return [serialize(child) for child in sorted(root["children"].values(), key=lambda item: item["name"])]


def _safe_extract_json(text: str) -> dict | None:
    if not text:
        return None
    cleaned = text.strip()
    cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s*```$", "", cleaned)
    try:
        data = json.loads(cleaned)
        return data if isinstance(data, dict) else None
    except json.JSONDecodeError:
        pass
    start = cleaned.find("{")
    end = cleaned.rfind("}")
    if start != -1 and end > start:
        try:
            data = json.loads(cleaned[start:end + 1])
            return data if isinstance(data, dict) else None
        except json.JSONDecodeError:
            return None
    return None


def _normalize_doc_type(value: str) -> str:
    value = (value or "").strip()
    if value in DOC_TYPES:
        return value
    aliases = {
        "problem": "题目",
        "problems": "题目",
        "question": "题目",
        "questions": "题目",
        "answer": "答案",
        "answers": "答案",
        "word": "词单",
        "words": "词单",
        "vocabulary": "词单",
        "material": "资料",
        "materials": "资料",
        "review": "资料",
    }
    return aliases.get(value.lower(), "资料")


def classify_material_locally(filename: str, text: str = "") -> dict:
    name = filename or ""
    lower = name.lower()
    sample = (text or "")[:3000]
    combined = f"{name}\n{sample}"

    doc_type = "资料"
    confidence = 0.55
    reason = "根据文件名和文本特征进行本地规则判断"

    word_keywords = ("单词", "词单", "词组", "默写")
    material_keywords = ("知识梳理", "知识点", "课文考点", "考点清单", "语言点", "讲义", "总结")
    answer_keywords = ("答案来源版", "答案版", "答案", "教师版", "解析", "批改")
    problem_keywords = ("学生版", "练习", "试卷", "课前测", "考题", "阅读", "完形", "填空")

    if any(key in name for key in word_keywords):
        doc_type = "词单"
        confidence = 0.82
        reason = "文件名包含单词/词单/默写等词单特征"
    elif any(key in name for key in material_keywords):
        doc_type = "资料"
        confidence = 0.82
        reason = "文件名包含知识梳理/考点/语言点等资料特征"
    elif any(key in name for key in answer_keywords):
        doc_type = "答案"
        confidence = 0.78
        reason = "文件名包含答案/教师版/解析等答案特征"
    elif re.search(r"【答案】|答案[:：]|answer\s*[:：]", combined, flags=re.IGNORECASE):
        doc_type = "答案"
        confidence = 0.72
        reason = "文本中出现答案标注"
    elif any(key in name for key in problem_keywords):
        doc_type = "题目"
        confidence = 0.72
        reason = "文件名包含学生版/练习/试卷等题目特征"
    elif re.search(r"\bA[.、)]\s+.+\bB[.、)]|\bC[.、)]|\bD[.、)]|_{3,}|____|c_{4,}", combined, flags=re.IGNORECASE | re.DOTALL):
        doc_type = "题目"
        confidence = 0.68
        reason = "文本中出现选择题选项或填空特征"
    elif re.search(r"\|.*(中文|英文|词性|pos|english|chinese).*\|", combined, flags=re.IGNORECASE):
        doc_type = "词单"
        confidence = 0.68
        reason = "文本中出现词单表格特征"

    summary = sample.strip().replace("\n", " ")[:120]
    return {
        "doc_type": doc_type,
        "confidence": confidence,
        "reason": reason,
        "summary": summary,
        "source": "local",
    }


def _find_material(materials: list[dict], material_id: str) -> tuple[int, dict] | tuple[None, None]:
    for idx, item in enumerate(materials):
        if item.get("id") == material_id:
            return idx, item
    return None, None


async def _load_current_user_materials(username: str) -> tuple[str, list[dict]]:
    config_subject = await get_user_subject(username)
    path = get_materials_path(username, config_subject)
    materials = await read_json(path) or []
    return path, materials if isinstance(materials, list) else []


def _cached_material_text(material_id: str) -> str:
    text_path = os.path.join(UPLOAD_DIR, f"{material_id}_text.txt")
    if os.path.exists(text_path):
        try:
            with open(text_path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception:
            return ""
    return ""


async def _validate_material_scope(user: dict, class_id: str, class_type_id: str) -> tuple[str, str]:
    class_id = (class_id or "").strip()
    class_type_id = (class_type_id or "").strip()
    if not class_id and not class_type_id:
        return "", ""

    classes = await load_classes()
    class_types = await load_class_types()
    if class_type_id and not any(item.get("id") == class_type_id for item in class_types):
        raise HTTPException(status_code=404, detail="班型不存在")

    username = user.get("username")
    role = user.get("role")
    if class_id:
        target_class = next((item for item in classes if item.get("id") == class_id), None)
        if target_class is None:
            raise HTTPException(status_code=404, detail="班级不存在")
        if role != "admin" and username not in (target_class.get("teacher_usernames") or []):
            raise HTTPException(status_code=403, detail="无权上传到该班级")
        if not class_type_id:
            class_type_id = target_class.get("class_type_id", "")

    if class_type_id and role != "admin":
        visible_type = any(
            item.get("class_type_id") == class_type_id and username in (item.get("teacher_usernames") or [])
            for item in classes
        )
        if not visible_type:
            raise HTTPException(status_code=403, detail="无权上传到该班型")

    return class_id, class_type_id


def _is_image(data: bytes) -> bool:
    return data.startswith(MAGIC_JPEG) or data.startswith(MAGIC_PNG)


def _is_pdf(data: bytes) -> bool:
    return data.startswith(MAGIC_PDF)


def _is_docx(data: bytes, filename: str = "") -> bool:
    if not data.startswith(MAGIC_ZIP):
        return False
    if filename.lower().endswith(".docx"):
        return True
    try:
        import zipfile
        import io
        with zipfile.ZipFile(io.BytesIO(data), "r") as z:
            return "word/document.xml" in z.namelist()
    except Exception:
        return False


def _run_to_html(run) -> str:
    text = run.text or ""
    if not text:
        return ""

    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = re.sub(r' {2,}', lambda m: '\xa0' * len(m.group()), text)

    font = run.font
    styles = []

    if font and font.size:
        try:
            styles.append(f"font-size:{font.size.pt}pt")
        except Exception:
            pass
    if font and font.name:
        styles.append(f"font-family:{font.name}")
    if font and font.color and font.color.rgb:
        styles.append(f"color:#{font.color.rgb}")

    style_attr = f' style="{";".join(styles)}"' if styles else ""

    if run.bold:
        text = f"<strong>{text}</strong>"
    if run.italic:
        text = f"<em>{text}</em>"
    if run.underline:
        text = f"<u>{text}</u>"
    if font and font.strike:
        text = f"<s>{text}</s>"
    if font and (font.subscript or font.superscript):
        tag = "sub" if font.subscript else "sup"
        text = f"<{tag}>{text}</{tag}>"

    if style_attr:
        text = f'<span{style_attr}>{text}</span>'

    text = text.replace("\n", "<br>").replace("\x0b", "<br>")
    return text


def _para_inner_html(para) -> str:
    return "".join(_run_to_html(run) for run in para.runs)


def _para_to_html(para) -> str:
    pf = para.paragraph_format
    styles = []

    if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        styles.append("text-align:center")
    elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        styles.append("text-align:right")
    elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        styles.append("text-align:justify")

    if pf.left_indent:
        try:
            styles.append(f"margin-left:{pf.left_indent.pt}pt")
        except Exception:
            pass
    if pf.right_indent:
        try:
            styles.append(f"margin-right:{pf.right_indent.pt}pt")
        except Exception:
            pass
    if pf.first_line_indent:
        try:
            styles.append(f"text-indent:{pf.first_line_indent.pt}pt")
        except Exception:
            pass
    if pf.space_before:
        try:
            styles.append(f"margin-top:{pf.space_before.pt}pt")
        except Exception:
            pass
    if pf.space_after:
        try:
            styles.append(f"margin-bottom:{pf.space_after.pt}pt")
        except Exception:
            pass
    if pf.line_spacing and isinstance(pf.line_spacing, (int, float)):
        styles.append(f"line-height:{pf.line_spacing}")

    ppr = para._p.get_or_add_pPr()
    br = ppr.find(qn('w:pageBreakBefore'))
    page_break = br is not None and br.get(qn('w:val')) != '0'

    style_attr = f' style="{";".join(styles)}"' if styles else ""

    heading_level = 0
    try:
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading "):
            heading_level = int(style_name.split()[-1])
    except Exception:
        pass

    content = _para_inner_html(para)
    if not content.strip():
        content = "&nbsp;"

    num_pr = ppr.find(qn('w:numPr'))
    is_list = num_pr is not None

    if is_list:
        return f"<li{style_attr}>{content}</li>"

    tag = f"h{heading_level}" if 1 <= heading_level <= 6 else "p"
    pb_attr = ' class="page-break"' if page_break else ""
    return f"<{tag}{pb_attr}{style_attr}>{content}</{tag}>"


def _table_to_html(table) -> str:
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_html = "".join(f"<p style=\"margin:2px 0;line-height:1.5;\">{_para_inner_html(p)}</p>" for p in cell.paragraphs if _para_inner_html(p).strip())
            cells.append(f'<td style="border:1px solid #bbb;padding:5px 8px;vertical-align:top;">{cell_html or "&nbsp;"}</td>')
        rows.append("<tr>" + "".join(cells) + "</tr>")
    return f'<table style="border-collapse:collapse;width:100%;margin:12px 0;font-size:10.5pt;">{"".join(rows)}</table>'


def _docx_to_html(data: bytes) -> str:
    try:
        doc = Document(io.BytesIO(data))
    except Exception:
        return ""

    html_parts = [
        '<div class="docx-content" style="font-family:Calibri,Microsoft YaHei,SimSun,sans-serif;font-size:11pt;line-height:1.7;color:#1e293b;">'
    ]

    para_idx = 0
    tbl_idx = 0
    list_open = None

    for element in doc.element.body:
        tag = element.tag.split('}')[-1]
        if tag == 'p':
            if para_idx < len(doc.paragraphs):
                para = doc.paragraphs[para_idx]
                para_html = _para_to_html(para)
                if para_html.startswith('<li'):
                    if list_open is None:
                        list_open = 'ul'
                        html_parts.append('<ul style="margin:4px 0;padding-left:28px;">')
                    html_parts.append(para_html)
                else:
                    if list_open is not None:
                        html_parts.append(f'</{list_open}>')
                        list_open = None
                    html_parts.append(para_html)
                para_idx += 1
        elif tag == 'tbl':
            if list_open is not None:
                html_parts.append(f'</{list_open}>')
                list_open = None
            if tbl_idx < len(doc.tables):
                html_parts.append(_table_to_html(doc.tables[tbl_idx]))
                tbl_idx += 1

    if list_open is not None:
        html_parts.append(f'</{list_open}>')

    html_parts.append('</div>')
    return "\n".join(html_parts)


async def _extract_docx_text(data: bytes) -> str:
    html = _docx_to_html(data)
    if not html:
        return ""
    text = re.sub(r'<br\s*/?>', '\n', html, flags=re.IGNORECASE)
    text = re.sub(r'</(p|div|li|tr|h[1-6]|td)>', '\n', text)
    text = re.sub(r'<[^>]+>', '', text)
    text = text.replace('\xa0', ' ')
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


async def _extract_pdf_text(data: bytes) -> str:
    try:
        import fitz
        tmp_id = str(uuid.uuid4())
        tmp_path = os.path.join(UPLOAD_DIR, f"{tmp_id}_extract.pdf")
        with open(tmp_path, "wb") as f:
            f.write(data)
        doc = fitz.open(tmp_path)
        parts = []
        for page in doc:
            try:
                parts.append(page.get_text("markdown"))
            except ValueError:
                parts.append(page.get_text())
        doc.close()
        try:
            os.unlink(tmp_path)
        except Exception:
            pass
        return "\n\n".join(parts).strip()
    except Exception:
        return ""


@router.post("")
async def upload_material(
    file: UploadFile = File(...),
    subject: str = Form(...),
    time: str = Form(...),
    tag: str = Form(""),
    doc_type: str = Form("资料"),
    class_id: str = Form(""),
    class_type_id: str = Form(""),
    user: dict = Depends(get_current_user_profile),
):
    username = user["username"]
    doc_type = (doc_type or "资料").strip()
    if doc_type not in DOC_TYPES:
        raise HTTPException(status_code=400, detail="资料类型无效")
    class_id, class_type_id = await _validate_material_scope(user, class_id, class_type_id)

    data = await file.read()
    if len(data) == 0:
        raise HTTPException(status_code=400, detail="文件内容为空")

    file_id = str(uuid.uuid4())
    file_type = "unknown"
    ext = ""
    filename = file.filename or ""

    if _is_pdf(data):
        file_type = "pdf"
        ext = ".pdf"
    elif _is_image(data):
        file_type = "image"
        ext = ".jpg"
    elif _is_docx(data, filename):
        file_type = "docx"
        ext = ".docx"
    else:
        raise HTTPException(status_code=400, detail="仅支持 PDF、Word（docx）、JPG、PNG 文件")

    save_path = os.path.join(UPLOAD_DIR, f"{file_id}{ext}")
    with open(save_path, "wb") as f:
        f.write(data)

    text_content = ""
    html_content = ""
    if file_type == "pdf":
        text_content = await _extract_pdf_text(data)
    elif file_type == "docx":
        html_content = _docx_to_html(data)
        text_content = await _extract_docx_text(data)

    if html_content and file_type == "docx":
        html_path = os.path.join(UPLOAD_DIR, f"{file_id}_text.html")
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(html_content)

    if text_content:
        text_path = os.path.join(UPLOAD_DIR, f"{file_id}_text.txt")
        with open(text_path, "w", encoding="utf-8") as f:
            f.write(text_content)

    material = {
        "id": file_id,
        "filename": filename or f"未命名{ext}",
        "subject": subject,
        "time": time,
        "tag": tag or time,
        "doc_type": doc_type,
        "owner_teacher": username,
        "class_id": class_id,
        "class_type_id": class_type_id,
        "file_path": save_path,
        "file_type": file_type,
        "file_size": len(data),
        "has_text": bool(text_content),
        "created_at": datetime.now().isoformat(),
    }

    config_subject = await get_user_subject(username)
    materials = await read_json(get_materials_path(username, config_subject)) or []
    materials.append(material)
    await write_json(get_materials_path(username, config_subject), materials)

    return material


@router.get("")
async def list_materials(
    subject: Optional[str] = None,
    time: Optional[str] = None,
    username: str = Depends(get_current_user),
):
    config_subject = await get_user_subject(username)
    materials = await read_json(get_materials_path(username, config_subject)) or []
    if subject:
        materials = [m for m in materials if m.get("subject") == subject]
    if time:
        materials = [m for m in materials if m.get("time") == time]
    return {"items": materials}


@router.get("/tree")
async def get_materials_tree(
    username: str = Depends(get_current_user),
):
    config_subject = await get_user_subject(username)
    materials = await read_json(get_materials_path(username, config_subject)) or []
    tree = {}
    for m in materials:
        tag = m.get("tag") or m.get("time", "未分类")
        if tag not in tree:
            tree[tag] = {"count": 0, "items": []}
        tree[tag]["count"] += 1
        tree[tag]["items"].append({
            "id": m["id"],
            "filename": m["filename"],
            "file_type": m.get("file_type", "pdf"),
            "has_text": m.get("has_text", False),
            "file_size": m.get("file_size", 0),
            "created_at": m.get("created_at", ""),
        })
    return tree


@router.get("/tag-tree")
async def get_materials_tag_tree(
    username: str = Depends(get_current_user),
):
    config_subject = await get_user_subject(username)
    materials = await read_json(get_materials_path(username, config_subject)) or []
    return {"items": _build_tag_tree(materials)}


@router.get("/tags")
async def get_tags(
    username: str = Depends(get_current_user),
):
    config_subject = await get_user_subject(username)
    materials = await read_json(get_materials_path(username, config_subject)) or []
    tags = sorted(set((m.get("tag") or m.get("time", "")) for m in materials if (m.get("tag") or m.get("time"))))
    return {"tags": tags}


@router.get("/tree_by_time")
async def get_materials_tree_by_time(
    username: str = Depends(get_current_user),
):
    config_subject = await get_user_subject(username)
    materials = await read_json(get_materials_path(username, config_subject)) or []
    tree = {}
    for m in materials:
        sub = m.get("subject", "未分类")
        t = m.get("time", "未分类")
        if sub not in tree:
            tree[sub] = {}
        if t not in tree[sub]:
            tree[sub][t] = {"count": 0, "items": []}
        tree[sub][t]["count"] += 1
        tree[sub][t]["items"].append({
            "id": m["id"],
            "filename": m["filename"],
            "file_type": m.get("file_type", "pdf"),
            "has_text": m.get("has_text", False),
        })
    return tree


@router.get("/subjects")
async def get_subjects(
    username: str = Depends(get_current_user),
):
    config_subject = await get_user_subject(username)
    materials = await read_json(get_materials_path(username, config_subject)) or []
    subjects = sorted(set(m.get("subject", "未分类") for m in materials if m.get("subject")))
    return {"subjects": subjects}


@router.get("/{material_id}")
async def get_material(
    material_id: str,
    username: str = Depends(get_current_user),
):
    config_subject = await get_user_subject(username)
    materials = await read_json(get_materials_path(username, config_subject)) or []
    for m in materials:
        if m["id"] == material_id:
            return m
    raise HTTPException(status_code=404, detail="资料不存在")


@router.post("/{material_id}/classify")
async def classify_material(
    material_id: str,
    user: dict = Depends(get_current_user_profile),
):
    username = user["username"]
    materials_path, materials = await _load_current_user_materials(username)
    idx, target = _find_material(materials, material_id)
    if target is None:
        raise HTTPException(status_code=404, detail="资料不存在")

    text = _cached_material_text(material_id)
    local_result = classify_material_locally(target.get("filename", ""), text)
    result = dict(local_result)

    config = await read_json(get_user_config_path(username)) or {}
    api_key = (config.get("kimi_api_key") or "").strip()
    if api_key:
        model = config.get("kimi_model", "kimi-k2.7-code")
        timeout = int(config.get("kimi_timeout", 120) or 120)
        prompt = f"""你是英语课外班资料分类助手。请根据文件名和文档内容，将资料分类为四类之一：题目、答案、词单、资料。

分类规则：
- 题目：学生练习、试卷、阅读/完形/选择/填空，含待作答空位。
- 答案：答案来源版、教师版答案、解析、批改讲解，含答案或参考解答。
- 词单：单词/词组/默写表，主要是英文、中文释义、词性对照。
- 资料：知识梳理、考点、语言点、讲义、语法总结。

只输出 JSON：
{{"doc_type":"题目|答案|词单|资料","confidence":0.0到1.0,"reason":"一句话理由","summary":"100字内摘要"}}

文件名：{target.get("filename", "")}
当前标签：{target.get("tag", "")}
文档内容节选：
{text[:5000] if text else "未提取到文本，请主要根据文件名判断"}
"""
        client = create_client(api_key, KIMI_BASE_URL, timeout)
        try:
            response = await client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "你是严格的英语教学资料分类器，只输出 JSON。"},
                    {"role": "user", "content": prompt},
                ],
                temperature=0.1,
                max_tokens=500,
                extra_body={"thinking": {"type": "disabled"}},
            )
            parsed = _safe_extract_json(response.choices[0].message.content or "")
            if parsed:
                result = {
                    "doc_type": _normalize_doc_type(parsed.get("doc_type") or parsed.get("type")),
                    "confidence": float(parsed.get("confidence", local_result["confidence"]) or local_result["confidence"]),
                    "reason": (parsed.get("reason") or local_result["reason"]).strip(),
                    "summary": (parsed.get("summary") or local_result["summary"]).strip(),
                    "source": "kimi",
                    "model": model,
                }
        except AuthenticationError:
            raise HTTPException(status_code=401, detail="Kimi API Key 无效")
        except RateLimitError:
            raise HTTPException(status_code=429, detail="Kimi 请求过于频繁")
        except APITimeoutError:
            result["reason"] += "；Kimi 超时，已使用本地规则结果"
        except (APIConnectionError, APIStatusError) as e:
            status = getattr(e, "status_code", 502)
            result["reason"] += f"；Kimi 返回错误 {status}，已使用本地规则结果"

    result["doc_type"] = _normalize_doc_type(result.get("doc_type", "资料"))
    result["confidence"] = max(0.0, min(1.0, float(result.get("confidence", 0.5) or 0.5)))
    target["doc_type"] = result["doc_type"]
    target["summary"] = result.get("summary", target.get("summary", ""))
    target["classification"] = {
        "doc_type": result["doc_type"],
        "confidence": result["confidence"],
        "reason": result.get("reason", ""),
        "source": result.get("source", "local"),
        "model": result.get("model", ""),
        "classified_at": datetime.now().isoformat(),
    }
    materials[idx] = target
    await write_json(materials_path, materials)
    return {"material": target, **target["classification"], "summary": target.get("summary", "")}


@router.patch("/{material_id}/type")
async def update_material_type(
    material_id: str,
    req: MaterialTypeUpdateRequest,
    user: dict = Depends(get_current_user_profile),
):
    doc_type = _normalize_doc_type(req.doc_type)
    if doc_type not in DOC_TYPES:
        raise HTTPException(status_code=400, detail="资料类型无效")
    materials_path, materials = await _load_current_user_materials(user["username"])
    idx, target = _find_material(materials, material_id)
    if target is None:
        raise HTTPException(status_code=404, detail="资料不存在")
    target["doc_type"] = doc_type
    target["manual_type_reason"] = req.reason.strip()
    target["manual_type_updated_at"] = datetime.now().isoformat()
    materials[idx] = target
    await write_json(materials_path, materials)
    return {"message": "资料类型已更新", "material": target}


@router.get("/{material_id}/text")
async def get_material_text(
    material_id: str,
    username: str = Depends(get_current_user),
):
    # Try cached _text.txt first
    text_path = os.path.join(UPLOAD_DIR, f"{material_id}_text.txt")
    if os.path.exists(text_path):
        with open(text_path, "r", encoding="utf-8") as f:
            content = f.read()
            if content.strip():
                return {"text": content}

    # Fallback: extract text from original file on-the-fly
    for ext in (".pdf", ".docx", ".jpg", ".jpeg", ".png"):
        candidate = os.path.join(UPLOAD_DIR, f"{material_id}{ext}")
        if not os.path.exists(candidate):
            continue
        try:
            if ext == ".pdf":
                text = await _extract_pdf_text(open(candidate, "rb").read())
            elif ext == ".docx":
                text = await _extract_docx_text(open(candidate, "rb").read())
            else:
                text = ""
            if text and text.strip():
                # Cache for next time
                try:
                    with open(text_path, "w", encoding="utf-8") as f:
                        f.write(text)
                except Exception:
                    pass
                return {"text": text}
        except Exception:
            continue
        break

    return {"text": ""}


@router.get("/{material_id}/html")
async def get_material_html(
    material_id: str,
    username: str = Depends(get_current_user),
):
    """Return rich HTML for docx files.

    1. Use already-cached ``{material_id}_text.html`` if present.
    2. Regenerate from original ``.docx`` on the fly and cache.
    3. Return empty — frontend will fall back to client-side mammoth.
    """
    html_path = os.path.join(UPLOAD_DIR, f"{material_id}_text.html")
    if os.path.exists(html_path):
        with open(html_path, "r", encoding="utf-8") as f:
            return {"html": f.read(), "type": "html"}

    candidate = os.path.join(UPLOAD_DIR, f"{material_id}.docx")
    if os.path.exists(candidate):
        try:
            with open(candidate, "rb") as f:
                data = f.read()
            html_content = _docx_to_html(data)
            if html_content:
                try:
                    with open(html_path, "w", encoding="utf-8") as f:
                        f.write(html_content)
                except Exception:
                    pass
                return {"html": html_content, "type": "html"}
        except Exception:
            pass

    # Fallback: show _text.txt content as plain pre-formatted text
    text_path = os.path.join(UPLOAD_DIR, f"{material_id}_text.txt")
    if os.path.exists(text_path):
        try:
            with open(text_path, "r", encoding="utf-8") as f:
                raw = f.read()
            # Strip common markdown formatting so users don't see raw syntax
            cleaned = raw.replace('***', '').replace('**', '').replace('*', '')
            cleaned = re.sub(r'^#+\s*', '', cleaned, flags=re.MULTILINE)
            cleaned = re.sub(r'^\s*\|[\s\-|]+\|\s*$', '', cleaned, flags=re.MULTILINE)
            cleaned = re.sub(r'\n{3,}', '\n\n', cleaned).strip()
            escaped = cleaned.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            return {
                "html": f'<div class="docx-content" style="white-space:pre-wrap;font-family:Calibri,Microsoft YaHei,SimSun,sans-serif;line-height:1.7;padding:20px;color:#334155;">{escaped}</div>',
                "type": "text",
            }
        except Exception:
            pass

    return {"html": "", "type": "none"}


@router.api_route("/{material_id}/file", methods=["GET", "HEAD"])
async def download_material_file(
    material_id: str,
    download: bool = False,
    username: str = Depends(get_current_user),
):
    config_subject = await get_user_subject(username)
    materials = await read_json(get_materials_path(username, config_subject)) or []
    target = None
    for m in materials:
        if m["id"] == material_id:
            target = m
            break

    file_path = ""
    filename = ""
    if target:
        file_path = target.get("file_path", "")
        filename = target.get("filename", "")

    if not file_path or not os.path.exists(file_path):
        for ext in (".pdf", ".docx", ".jpg", ".jpeg", ".png"):
            candidate = os.path.join(UPLOAD_DIR, f"{material_id}{ext}")
            if os.path.exists(candidate):
                file_path = candidate
                if not filename:
                    filename = os.path.basename(candidate)
                break

    if not file_path or not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="原始文件不存在")

    if not filename:
        filename = os.path.basename(file_path)

    return await _serve_file(file_path, filename, download)


@router.get("/{material_id}/pdf")
async def get_material_pdf(
    material_id: str,
    username: str = Depends(get_current_user),
):
    """Convert docx to PDF using LibreOffice and serve the PDF inline."""
    config_subject = await get_user_subject(username)
    materials = await read_json(get_materials_path(username, config_subject)) or []
    target = None
    for m in materials:
        if m["id"] == material_id:
            target = m
            break

    if not target:
        raise HTTPException(status_code=404, detail="资料不存在")

    file_type = (target.get("file_type") or "").lower()
    if file_type != "docx":
        if file_type == "pdf":
            file_path = target.get("file_path", "")
            if not file_path or not os.path.exists(file_path):
                for ext in (".pdf",):
                    candidate = os.path.join(UPLOAD_DIR, f"{material_id}{ext}")
                    if os.path.exists(candidate):
                        file_path = candidate
                        break
            if file_path and os.path.exists(file_path):
                return await _serve_file(file_path, target.get("filename", ""), False)
        raise HTTPException(status_code=400, detail="仅支持 docx 文件转换 PDF")

    from app.core.paths import BASE_DIR

    docx_path = ""
    docx_path = target.get("file_path", "")
    if not docx_path or not os.path.exists(docx_path):
        candidate = os.path.join(UPLOAD_DIR, f"{material_id}.docx")
        if os.path.exists(candidate):
            docx_path = candidate

    # Try resolving relative paths against BASE_DIR
    if not docx_path or not os.path.exists(docx_path):
        relative = target.get("file_path", "")
        if relative and not os.path.isabs(relative):
            abs_path = os.path.normpath(os.path.join(BASE_DIR, relative))
            if os.path.exists(abs_path):
                docx_path = abs_path

    if not docx_path or not os.path.exists(docx_path):
        raise HTTPException(status_code=404, detail="原始文件不存在")

    logger = logging.getLogger("materials")
    logger.info(f"Converting to PDF: docx_path={docx_path} UPLOAD_DIR={UPLOAD_DIR}")
    pdf_path = await convert_docx_to_pdf(docx_path, UPLOAD_DIR)
    if not pdf_path or not os.path.exists(pdf_path):
        logger.error(f"PDF conversion failed: docx_path={docx_path} UPLOAD_DIR={UPLOAD_DIR}")
        from app.utils.doc_converter import _soffice_exists, SOFFICE
        logger.error(f"soffice_exists={_soffice_exists()} SOFFICE={SOFFICE}")
        # Try sync fallback
        pdf_path = _convert_sync(docx_path, UPLOAD_DIR)
        if not pdf_path or not os.path.exists(pdf_path):
            raise HTTPException(status_code=500, detail="PDF 转换失败，请检查 LibreOffice 是否正常安装")

    pdf_filename = os.path.splitext(target.get("filename", "document"))[0] + ".pdf"
    return await _serve_file(pdf_path, pdf_filename, False)


@router.get("/{material_id}/print")
async def print_material(
    material_id: str,
    username: str = Depends(get_current_user),
):
    """Return an HTML page that embeds the PDF and auto-prints once loaded."""
    config_subject = await get_user_subject(username)
    materials = await read_json(get_materials_path(username, config_subject)) or []
    target = None
    for m in materials:
        if m["id"] == material_id:
            target = m
            break

    if not target:
        raise HTTPException(status_code=404, detail="资料不存在")

    file_type = (target.get("file_type") or "").lower()
    pdf_id = material_id
    if file_type == "docx":
        # Trigger conversion first
        from app.core.paths import BASE_DIR
        docx_path = target.get("file_path", "")
        if not docx_path or not os.path.exists(docx_path):
            candidate = os.path.join(UPLOAD_DIR, f"{material_id}.docx")
            if os.path.exists(candidate):
                docx_path = candidate
        if docx_path and os.path.exists(docx_path):
            await convert_docx_to_pdf(docx_path, UPLOAD_DIR)

    pdf_url = f"/api/materials/{material_id}/pdf"
    html = f"""<!DOCTYPE html>
<html><head>
<meta charset="utf-8">
<style>
body {{ margin: 0; }}
iframe {{ width: 100vw; height: 100vh; border: none; }}
@media print {{ .no-print {{ display: none; }} }}
</style>
</head><body>
<div id="toolbar" style="position:fixed;top:0;right:0;z-index:999;padding:10px;background:rgba(255,255,255,0.9);">
<button onclick="window.print()" style="padding:8px 20px;font-size:14px;cursor:pointer;border:1px solid #ccc;border-radius:4px;background:#2563eb;color:#fff;">打印</button>
</div>
<iframe id="pdf-frame" src="{pdf_url}" loading="lazy"></iframe>
<script>
var iframe = document.getElementById('pdf-frame');
if (iframe) {{
  var attemptPrint = function(attempts) {{
    attempts = attempts || 0;
    if (attempts > 20) return;
    try {{
      window.print();
    }} catch(e) {{
      setTimeout(function() {{ attemptPrint(attempts + 1); }}, 500);
    }}
  }};
  iframe.onload = function() {{
    setTimeout(function() {{ attemptPrint(); }}, 1500);
  }};
}}
</script>
</body></html>"""
    return HTMLResponse(content=html)


async def _serve_file(file_path: str, filename: str, download: bool):
    """Helper to serve a file with proper content type and disposition."""
    ext = os.path.splitext(file_path)[1].lower()
    media_type = "application/octet-stream"
    if ext == ".pdf":
        media_type = "application/pdf"
    elif ext == ".docx":
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif ext in (".jpg", ".jpeg"):
        media_type = "image/jpeg"
    elif ext == ".png":
        media_type = "image/png"

    disposition_type = "attachment" if download else "inline"
    encoded_name = quote(filename, safe="")
    ascii_safe = ''.join(c if 32 <= ord(c) < 127 else '_' for c in filename)
    ascii_safe = ascii_safe.replace('\\', '\\\\').replace('"', '\\"')
    cd_header = f'{disposition_type}; filename="{ascii_safe}"; filename*=UTF-8\'\'{encoded_name}'

    headers = {
        "Content-Disposition": cd_header,
        "Content-Length": str(os.path.getsize(file_path)),
    }

    def file_iterator():
        with open(file_path, "rb") as f:
            while True:
                chunk = f.read(1024 * 1024)
                if not chunk:
                    break
                yield chunk

    return StreamingResponse(
        file_iterator(),
        media_type=media_type,
        headers=headers,
    )


@router.delete("/{material_id}")
async def delete_material(
    material_id: str,
    username: str = Depends(get_current_user),
):
    config_subject = await get_user_subject(username)
    materials = await read_json(get_materials_path(username, config_subject)) or []
    target = None
    for m in materials:
        if m["id"] == material_id:
            target = m
            break
    if not target:
        raise HTTPException(status_code=404, detail="资料不存在")
    target["trashed_at"] = datetime.now().isoformat()
    target["_origin_type"] = "material"
    new_materials = [x for x in materials if x["id"] != material_id]
    await write_json(get_materials_path(username, config_subject), new_materials)
    trash = await read_json(get_library_trash_path(username, config_subject)) or []
    trash.append(target)
    await write_json(get_library_trash_path(username, config_subject), trash)
    return {"message": "已移入废纸篓"}
