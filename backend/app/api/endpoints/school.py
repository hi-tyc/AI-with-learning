from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Literal, Optional
from openai import APITimeoutError, APIConnectionError, APIStatusError, RateLimitError, AuthenticationError
import base64
import csv
import io
import json
import os
import re
import zipfile
from urllib.parse import quote

from app.api.endpoints.auth import (
    _default_config,
    _default_user,
    _hash_password,
    get_current_user_profile,
    require_roles,
)
from app.core.paths import REGISTRATION_DIR, UPLOAD_DIR, USERS_DIR
from app.core.school_data import (
    load_class_types,
    save_class_types,
    load_classes,
    save_classes,
    load_distributions,
    save_distributions,
    new_entity_id,
    now_iso,
    normalize_students,
    parse_roster_text,
    sync_student_to_classes,
)
from app.core.user_data import get_materials_path, get_user_config_path, get_user_path, get_user_subject
from app.utils.file_lock import read_json, write_json
from app.utils.ai_client import create_client, KIMI_BASE_URL

router = APIRouter()

AUTO_GRADE_CONFIDENCE_THRESHOLD = 0.75


class ClassTypeCreateRequest(BaseModel):
    name: str
    description: str = ""


class ClassCreateRequest(BaseModel):
    name: str
    class_type_id: str
    description: str = ""
    teacher_usernames: list[str] = []
    roster_text: str = ""


class ManagedUserCreateRequest(BaseModel):
    username: str
    real_name: str
    role: Literal["teacher", "student"] = "teacher"
    password: str = ""
    class_ids: list[str] = []
    class_type_ids: list[str] = []
    expires_at: Optional[str] = None
    activate_now: bool = True


class DistributionCreateRequest(BaseModel):
    material_ids: list[str]
    target_type: Literal["class_type", "class"]
    target_ids: list[str]
    tag_path: str = ""
    note: str = ""


class CompletionMarkRequest(BaseModel):
    student_name: str
    completed_parts: list[str] = []
    note: str = ""


class WrongBookMarkRequest(BaseModel):
    student_name: str
    problem_ref: str
    is_wrong: bool = True
    note: str = ""


class ExplanationCreateRequest(BaseModel):
    problem_ref: str
    force: bool = False


class AutoGradeRequest(BaseModel):
    student_name: str
    answer_material_id: str
    submission_material_id: str
    replace_existing: bool = True


class RegistrationReviewRequest(BaseModel):
    action: Literal["approve", "reject"]
    class_ids: list[str] = []
    expires_at: Optional[str] = None
    note: str = ""


def _dedupe(items: list[str]) -> list[str]:
    return list(dict.fromkeys([item.strip() for item in items if item and item.strip()]))


def _find_by_id(items: list[dict], entity_id: str) -> dict | None:
    for item in items:
        if item.get("id") == entity_id:
            return item
    return None


def _class_visible_to_user(class_item: dict, user: dict) -> bool:
    if user.get("role") == "admin":
        return True
    return user.get("username") in (class_item.get("teacher_usernames") or [])


def _distribution_visible_to_user(item: dict, user: dict) -> bool:
    if user.get("role") == "admin":
        return True
    return item.get("assigned_by") == user.get("username")


def _compute_distribution_stats(item: dict) -> dict:
    students = item.get("students") or []
    student_total = len(students)
    wrong_records = [row for row in (item.get("wrong_book_records") or []) if row.get("is_wrong")]
    completion_records = item.get("completion_records") or []
    auto_grade_records = item.get("auto_grade_records") or []

    wrong_by_problem: dict[str, list[str]] = {}
    for row in wrong_records:
        problem_ref = row.get("problem_ref") or "未命名题目"
        wrong_by_problem.setdefault(problem_ref, [])
        student_name = row.get("student_name") or ""
        if student_name and student_name not in wrong_by_problem[problem_ref]:
            wrong_by_problem[problem_ref].append(student_name)

    problem_stats = []
    explanations = item.get("explanations") or {}
    for problem_ref, student_names in sorted(wrong_by_problem.items()):
        wrong_count = len(student_names)
        correct_rate = 0.0
        if student_total > 0:
            correct_rate = max(0.0, (student_total - wrong_count) / student_total)
        explanation = explanations.get(problem_ref) if isinstance(explanations, dict) else None
        problem_stats.append({
            "problem_ref": problem_ref,
            "wrong_students": student_names,
            "wrong_count": wrong_count,
            "correct_rate": round(correct_rate, 4),
            "explanation": explanation.get("content", "") if isinstance(explanation, dict) else "",
            "explanation_model": explanation.get("model", "") if isinstance(explanation, dict) else "",
            "explanation_generated_at": explanation.get("generated_at") if isinstance(explanation, dict) else None,
        })

    return {
        "student_total": student_total,
        "wrong_count": len(wrong_records),
        "completion_count": len(completion_records),
        "auto_grade_count": len(auto_grade_records),
        "problem_stats": problem_stats,
    }


def _csv_response(filename: str, headers: list[str], rows: list[list[str | int | float]]) -> StreamingResponse:
    buffer = io.StringIO()
    writer = csv.writer(buffer)
    writer.writerow(headers)
    for row in rows:
        writer.writerow(row)
    payload = ("\ufeff" + buffer.getvalue()).encode("utf-8")
    return StreamingResponse(
        iter([payload]),
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _collect_wrong_book_entries(items: list[dict], student_name: str = "") -> list[dict]:
    entries_by_key: dict[tuple[str, str, str], dict] = {}
    for item in items:
        stats = _compute_distribution_stats(item)
        stat_map = {row["problem_ref"]: row for row in stats["problem_stats"]}
        explanations = item.get("explanations") or {}
        for row in item.get("wrong_book_records") or []:
            if not row.get("is_wrong"):
                continue
            row_student = row.get("student_name", "")
            if student_name and row_student != student_name:
                continue

            problem_ref = row.get("problem_ref", "") or "未命名题目"
            key_student = row_student if student_name else ""
            key = (item.get("id", ""), problem_ref, key_student)
            stat = stat_map.get(problem_ref, {})
            explanation = explanations.get(problem_ref) if isinstance(explanations, dict) else None
            entry = entries_by_key.setdefault(key, {
                "distribution_id": item.get("id", ""),
                "problem_ref": problem_ref,
                "students": [],
                "notes": [],
                "explanation": explanation.get("content", "") if isinstance(explanation, dict) else "",
                "wrong_students": stat.get("wrong_students", []),
                "correct_rate": stat.get("correct_rate"),
                "wrong_count": stat.get("wrong_count"),
                "student_total": stats.get("student_total", 0),
                "tag_path": item.get("tag_path", ""),
                "distribution_note": item.get("note", ""),
                "assigned_at": item.get("assigned_at", ""),
            })
            if row_student and row_student not in entry["students"]:
                entry["students"].append(row_student)
            note = (row.get("note") or "").strip()
            if note:
                note_text = f"{row_student}: {note}" if row_student and not student_name else note
                if note_text not in entry["notes"]:
                    entry["notes"].append(note_text)

    return sorted(
        entries_by_key.values(),
        key=lambda item: (item.get("assigned_at", ""), item.get("tag_path", ""), item.get("problem_ref", "")),
    )


def _wrong_book_docx_zip_response(entries: list[dict], title: str, filename: str) -> StreamingResponse:
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(status_code=500, detail="缺少 python-docx，无法生成 Word 文档")

    def setup_doc(doc: Document, subtitle: str) -> None:
        section = doc.sections[0]
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        doc.styles["Normal"].font.size = Pt(10.5)
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.runs[0].font.size = Pt(16)
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = meta.add_run(f"{subtitle}  |  共 {len(entries)} 题  |  生成时间：{now_iso()[:16].replace('T', ' ')}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    def add_meta_line(doc: Document, label: str, text: str) -> None:
        if not text:
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r_label = p.add_run(label)
        r_label.bold = True
        r_label.font.size = Pt(9)
        r_text = p.add_run(text)
        r_text.font.size = Pt(9)

    student_doc = Document()
    answer_doc = Document()
    setup_doc(student_doc, "学生版")
    setup_doc(answer_doc, "答案版")

    if not entries:
        student_doc.add_paragraph("暂无错题记录。")
        answer_doc.add_paragraph("暂无错题记录。")

    for index, entry in enumerate(entries, start=1):
        problem_ref = entry.get("problem_ref", "")
        for doc in (student_doc, answer_doc):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            r_num = p.add_run(f"{index}. ")
            r_num.bold = True
            r_num.font.size = Pt(11)
            r_body = p.add_run(problem_ref)
            r_body.font.size = Pt(11)
            add_meta_line(doc, "标签：", entry.get("tag_path", ""))
            add_meta_line(doc, "分发备注：", entry.get("distribution_note", ""))

        answer_parts = []
        if entry.get("students"):
            answer_parts.append("涉及学生：" + "、".join(entry["students"]))
        if entry.get("wrong_students"):
            answer_parts.append("错误学生：" + "、".join(entry["wrong_students"]))
        if entry.get("correct_rate") is not None:
            answer_parts.append(f"正确率：{entry['correct_rate'] * 100:.2f}%")
        if entry.get("notes"):
            answer_parts.append("备注：" + "；".join(entry["notes"]))
        if entry.get("explanation"):
            answer_parts.append("AI解析：" + entry["explanation"])

        if answer_parts:
            for part in answer_parts:
                para = answer_doc.add_paragraph()
                para.paragraph_format.space_after = Pt(2)
                run = para.add_run(part)
                run.font.size = Pt(10.5)
        else:
            answer_doc.add_paragraph("暂无备注或解析。")

    student_buf = io.BytesIO()
    answer_buf = io.BytesIO()
    student_doc.save(student_buf)
    answer_doc.save(answer_buf)
    student_buf.seek(0)
    answer_buf.seek(0)

    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"{title}_学生版.docx", student_buf.getvalue())
        zf.writestr(f"{title}_答案版.docx", answer_buf.getvalue())
    zip_buf.seek(0)

    encoded_name = quote(filename, safe="")
    ascii_name = "".join(ch if 32 <= ord(ch) < 127 else "_" for ch in filename)
    headers = {"Content-Disposition": f"attachment; filename=\"{ascii_name}\"; filename*=UTF-8''{encoded_name}"}
    return StreamingResponse(iter([zip_buf.getvalue()]), media_type="application/zip", headers=headers)


def _filter_visible_distributions(items: list[dict], user: dict) -> list[dict]:
    if user.get("role") == "admin":
        return items
    return [item for item in items if _distribution_visible_to_user(item, user)]


def _delete_registration_video(registration: dict, reason: str) -> dict:
    video_path = registration.get("face_video_path", "")
    if video_path:
        abs_video_path = os.path.abspath(video_path)
        abs_registration_dir = os.path.abspath(REGISTRATION_DIR)
        if abs_video_path.startswith(abs_registration_dir + os.sep) and os.path.exists(abs_video_path):
            os.remove(abs_video_path)
    registration["face_video_name"] = ""
    registration["face_video_path"] = ""
    registration["face_video_deleted_at"] = now_iso()
    registration["face_video_delete_reason"] = reason
    return registration


def _wrong_book_context(item: dict, problem_ref: str) -> dict:
    rows = [
        row for row in (item.get("wrong_book_records") or [])
        if row.get("is_wrong") and (row.get("problem_ref") or "") == problem_ref
    ]
    wrong_students = []
    notes = []
    for row in rows:
        student_name = row.get("student_name") or ""
        if student_name and student_name not in wrong_students:
            wrong_students.append(student_name)
        note = (row.get("note") or "").strip()
        if note:
            notes.append(f"{student_name}: {note}" if student_name else note)
    return {
        "wrong_students": wrong_students,
        "notes": notes,
        "student_total": len(item.get("students") or []),
        "distribution_note": item.get("note", ""),
        "tag_path": item.get("tag_path", ""),
    }


async def _generate_kimi_explanation(problem_ref: str, context: dict, user: dict) -> dict:
    config = await read_json(get_user_config_path(user["username"])) or {}
    api_key = config.get("kimi_api_key", "")
    if not api_key:
        raise HTTPException(status_code=400, detail="Kimi API Key 未配置")
    model = config.get("kimi_model", "kimi-k2.7-code")
    timeout = int(config.get("kimi_timeout", 120) or 120)
    prompt = f"""你是英语课外班老师的助教。请基于下面的错题信息，生成给学生看的简洁解析。

要求：
1. 只输出中文解析，不要编造不存在的原题细节。
2. 如果题目标识不是完整题干，请先说明“题干信息有限”，再给出通用解题思路。
3. 包含：考点、易错点、解题步骤、订正建议。
4. 控制在 300 字以内。

题目标识或题干：
{problem_ref}

错误学生：{"、".join(context["wrong_students"]) or "未记录"}
学生总数：{context["student_total"]}
标签路径：{context["tag_path"] or "未分类"}
分发备注：{context["distribution_note"] or "无"}
错题备注：
{chr(10).join(context["notes"]) or "无"}
"""
    client = create_client(api_key, KIMI_BASE_URL, timeout)
    try:
        response = await client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "你是严谨的英语学习错题解析助手。"},
                {"role": "user", "content": prompt},
            ],
            temperature=0.2,
            max_tokens=700,
            extra_body={"thinking": {"type": "disabled"}},
        )
        content = (response.choices[0].message.content or "").strip()
        if not content:
            raise HTTPException(status_code=502, detail="Kimi 未返回解析内容")
        return {
            "content": content,
            "model": model,
            "generated_by": user["username"],
            "generated_at": now_iso(),
        }
    except AuthenticationError:
        raise HTTPException(status_code=401, detail="Kimi API Key 无效")
    except RateLimitError:
        raise HTTPException(status_code=429, detail="Kimi 请求过于频繁")
    except APITimeoutError:
        raise HTTPException(status_code=504, detail="Kimi 解析生成超时")
    except (APIConnectionError, APIStatusError) as e:
        status = getattr(e, "status_code", 502)
        raise HTTPException(status_code=status, detail=f"Kimi 返回错误: {status}")


def _safe_json_object(text: str) -> dict | None:
    """Extract one JSON object from a Kimi response without trusting surrounding prose."""
    if not text:
        return None
    cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", text.strip(), flags=re.IGNORECASE)
    try:
        payload = json.loads(cleaned)
        return payload if isinstance(payload, dict) else None
    except json.JSONDecodeError:
        pass
    start = cleaned.find("{")
    end = cleaned.rfind("}")
    if start < 0 or end <= start:
        return None
    try:
        payload = json.loads(cleaned[start:end + 1])
        return payload if isinstance(payload, dict) else None
    except json.JSONDecodeError:
        return None


def _is_wrong_value(value) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return bool(value)
    normalized = str(value or "").strip().lower()
    return normalized in {"true", "1", "wrong", "incorrect", "error", "错误", "错", "不正确"}


def _normalize_auto_grade_result(content: str) -> dict:
    payload = _safe_json_object(content)
    if payload is None:
        raise HTTPException(status_code=502, detail="Kimi 批改结果格式无效，请重新批改")

    raw_items = payload.get("items") or payload.get("results") or payload.get("wrong_items") or []
    if not isinstance(raw_items, list):
        raise HTTPException(status_code=502, detail="Kimi 批改结果缺少题目明细，请重新批改")

    items = []
    for index, raw in enumerate(raw_items, start=1):
        if not isinstance(raw, dict):
            continue
        problem_ref = str(raw.get("problem_ref") or raw.get("question") or raw.get("question_no") or f"第{index}题").strip()
        if not problem_ref:
            problem_ref = f"第{index}题"
        student_answer = str(raw.get("student_answer") or raw.get("answer") or "").strip()
        correct_answer = str(raw.get("correct_answer") or raw.get("reference_answer") or "").strip()
        reason = str(raw.get("reason") or raw.get("comment") or raw.get("feedback") or "").strip()
        confidence_raw = raw.get("confidence", 0)
        try:
            confidence = max(0.0, min(1.0, float(confidence_raw)))
        except (TypeError, ValueError):
            confidence = 0.0
        items.append({
            "problem_ref": problem_ref,
            "student_answer": student_answer,
            "correct_answer": correct_answer,
            "is_wrong": _is_wrong_value(raw.get("is_wrong", raw.get("wrong", raw.get("result")))),
            "reason": reason,
            "confidence": round(confidence, 4),
        })

    score_raw = payload.get("score")
    try:
        score = max(0.0, min(100.0, float(score_raw))) if score_raw is not None else None
    except (TypeError, ValueError):
        score = None
    summary = str(payload.get("summary") or payload.get("overall_comment") or "").strip()
    candidate_wrong_items = [item for item in items if item["is_wrong"]]
    wrong_items = [
        item for item in candidate_wrong_items
        if item["confidence"] >= AUTO_GRADE_CONFIDENCE_THRESHOLD
    ]
    try:
        reported_ungraded_count = max(0, int(float(payload.get("ungraded_count", 0) or 0)))
    except (TypeError, ValueError):
        reported_ungraded_count = 0
    return {
        "summary": summary,
        "score": score,
        "items": items,
        "wrong_items": wrong_items,
        "recognized_count": len(items),
        "wrong_count": len(wrong_items),
        "ungraded_count": max(reported_ungraded_count, len(candidate_wrong_items) - len(wrong_items)),
    }


def _material_cached_text(material: dict) -> str:
    material_id = str(material.get("id") or "").strip()
    if not material_id:
        return ""
    text_path = os.path.join(UPLOAD_DIR, f"{material_id}_text.txt")
    try:
        with open(text_path, "r", encoding="utf-8") as handle:
            return handle.read()
    except OSError:
        return ""


def _material_image_data_url(material: dict) -> str:
    if material.get("file_type") != "image":
        return ""
    path = os.path.abspath(str(material.get("file_path") or ""))
    upload_root = os.path.abspath(UPLOAD_DIR)
    if not path or os.path.commonpath([upload_root, path]) != upload_root:
        return ""
    try:
        if os.path.getsize(path) > 8 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="用于 AI 批改的图片不能超过 8MB")
        with open(path, "rb") as handle:
            data = handle.read()
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail=f"找不到资料文件：{material.get('filename') or material.get('id')}")
    except OSError:
        return ""
    if not data:
        return ""
    mime_type = "image/png" if data.startswith(b"\x89PNG\r\n\x1a\n") else "image/jpeg"
    encoded = base64.b64encode(data).decode("ascii")
    return f"data:{mime_type};base64,{encoded}"


def _material_grade_parts(label: str, material: dict) -> list[dict]:
    text = _material_cached_text(material).strip()
    image_data_url = _material_image_data_url(material)
    if not text and not image_data_url:
        raise HTTPException(
            status_code=422,
            detail=f"{label}没有可识别的文本或图片，请上传清晰的 JPG、PNG、PDF 或 Word 文件",
        )
    parts = [{
        "type": "text",
        "text": f"{label}文件：{material.get('filename') or material.get('id')}\n"
                f"{label}提取文本：\n{text[:24000] or '（无文本，见附图）'}",
    }]
    if image_data_url:
        parts.append({"type": "image_url", "image_url": {"url": image_data_url}})
    return parts


async def _run_kimi_auto_grade(answer_material: dict, submission_material: dict, user: dict) -> dict:
    config = await read_json(get_user_config_path(user["username"])) or {}
    api_key = config.get("kimi_api_key", "")
    if not api_key:
        raise HTTPException(status_code=400, detail="Kimi API Key 未配置")
    model = config.get("kimi_model", "kimi-k2.7-code")
    timeout = int(config.get("kimi_timeout", 120) or 120)
    prompt = """你是严谨的英语教师。请比较“标准答案”和“学生作答”，只依据提供的内容批改，不要猜测看不清的信息。

请只返回一个 JSON 对象，不要 Markdown，不要额外解释。格式必须是：
{
  "summary": "给学生的整体简短反馈",
  "score": 0 到 100 的数字或 null,
  "ungraded_count": 未能可靠判断的题数,
  "items": [
    {
      "problem_ref": "第1题或题干简写",
      "student_answer": "学生答案",
      "correct_answer": "标准答案",
      "is_wrong": true 或 false,
      "reason": "错因或判定依据",
      "confidence": 0 到 1 的数字
    }
  ]
}

规则：
1. 只有能够明确比对的错误才设置 is_wrong 为 true。
2. 识别不清、题号无法确认时不要标错，计入 ungraded_count。
3. 题号尽量与标准答案一致，避免把多个小题合并。
4. 避免把格式、大小写等无关差异判成错误，除非标准答案要求明确。
"""
    content = [{"type": "text", "text": prompt}]
    content.extend(_material_grade_parts("标准答案", answer_material))
    content.extend(_material_grade_parts("学生作答", submission_material))
    client = create_client(api_key, KIMI_BASE_URL, timeout)
    try:
        response = await client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "你只做依据充分的英语作业批改，并严格输出 JSON。"},
                {"role": "user", "content": content},
            ],
            temperature=0.1,
            max_tokens=3500,
            extra_body={"thinking": {"type": "disabled"}},
        )
        result_text = (response.choices[0].message.content or "").strip()
        if not result_text:
            raise HTTPException(status_code=502, detail="Kimi 未返回批改结果")
        result = _normalize_auto_grade_result(result_text)
        result["model"] = model
        return result
    except AuthenticationError:
        raise HTTPException(status_code=401, detail="Kimi API Key 无效")
    except RateLimitError:
        raise HTTPException(status_code=429, detail="Kimi 请求过于频繁")
    except APITimeoutError:
        raise HTTPException(status_code=504, detail="Kimi 自动批改超时")
    except (APIConnectionError, APIStatusError) as error:
        status = getattr(error, "status_code", 502)
        raise HTTPException(status_code=status, detail=f"Kimi 返回错误: {status}")


def _apply_auto_grade_result(target: dict, req: AutoGradeRequest, result: dict, user: dict) -> tuple[dict, int]:
    student_name = req.student_name.strip()
    answer_material_id = req.answer_material_id.strip()
    submission_material_id = req.submission_material_id.strip()
    wrong_records = target.setdefault("wrong_book_records", [])
    if req.replace_existing:
        wrong_records[:] = [
            row for row in wrong_records
            if not (
                row.get("source") == "ai_auto_grade"
                and row.get("student_name") == student_name
                and row.get("answer_material_id") == answer_material_id
                and row.get("submission_material_id") == submission_material_id
            )
        ]

    grade_id = new_entity_id("grade")
    created_count = 0
    for item in result["wrong_items"]:
        note_parts = ["AI 自动批改"]
        if item.get("student_answer"):
            note_parts.append(f"学生答案：{item['student_answer']}")
        if item.get("correct_answer"):
            note_parts.append(f"标准答案：{item['correct_answer']}")
        if item.get("reason"):
            note_parts.append(f"判定：{item['reason']}")
        wrong_records.append({
            "student_name": student_name,
            "problem_ref": item["problem_ref"],
            "is_wrong": True,
            "note": "；".join(note_parts)[:1600],
            "source": "ai_auto_grade",
            "auto_grade_id": grade_id,
            "answer_material_id": answer_material_id,
            "submission_material_id": submission_material_id,
            "confidence": item.get("confidence", 0),
            "marked_by": user["username"],
            "marked_at": now_iso(),
        })
        created_count += 1

    grade_record = {
        "id": grade_id,
        "student_name": student_name,
        "answer_material_id": answer_material_id,
        "submission_material_id": submission_material_id,
        "summary": result.get("summary", ""),
        "score": result.get("score"),
        "recognized_count": result.get("recognized_count", 0),
        "wrong_count": result.get("wrong_count", 0),
        "ungraded_count": result.get("ungraded_count", 0),
        "wrong_items": result.get("wrong_items", []),
        "model": result.get("model", ""),
        "graded_by": user["username"],
        "graded_at": now_iso(),
    }
    target.setdefault("auto_grade_records", []).append(grade_record)
    return grade_record, created_count


@router.get("/summary")
async def school_summary(user: dict = Depends(get_current_user_profile)):
    class_types = await load_class_types()
    classes = await load_classes()
    distributions = await load_distributions()
    if user.get("role") != "admin":
        classes = [item for item in classes if _class_visible_to_user(item, user)]
        distributions = [item for item in distributions if _distribution_visible_to_user(item, user)]
    student_total = sum(len(item.get("students") or []) for item in classes)
    return {
        "class_type_count": len(class_types),
        "class_count": len(classes),
        "distribution_count": len(distributions),
        "student_total": student_total,
    }


@router.get("/class-types")
async def list_class_types(user: dict = Depends(get_current_user_profile)):
    return {"items": await load_class_types()}


@router.get("/public-classes")
async def list_public_classes():
    class_types = await load_class_types()
    class_type_map = {item["id"]: item for item in class_types}
    items = await load_classes()
    return {
        "items": [
            {
                "id": item.get("id"),
                "name": item.get("name"),
                "class_type_id": item.get("class_type_id"),
                "class_type_name": class_type_map.get(item.get("class_type_id"), {}).get("name", ""),
            }
            for item in items
        ]
    }


@router.post("/class-types")
async def create_class_type(
    req: ClassTypeCreateRequest,
    user: dict = Depends(require_roles("admin")),
):
    name = req.name.strip()
    if not name:
        raise HTTPException(status_code=400, detail="班型名称不能为空")
    items = await load_class_types()
    if any(item.get("name") == name for item in items):
        raise HTTPException(status_code=409, detail="班型已存在")
    item = {
        "id": new_entity_id("ctype"),
        "name": name,
        "description": req.description.strip(),
        "created_at": now_iso(),
        "created_by": user["username"],
    }
    items.append(item)
    await save_class_types(items)
    return item


@router.get("/classes")
async def list_classes(user: dict = Depends(get_current_user_profile)):
    class_types = await load_class_types()
    class_type_map = {item["id"]: item for item in class_types}
    items = await load_classes()
    if user.get("role") != "admin":
        items = [item for item in items if _class_visible_to_user(item, user)]
    enriched = []
    for item in items:
        row = dict(item)
        row["class_type_name"] = class_type_map.get(item.get("class_type_id"), {}).get("name", "")
        row["student_count"] = len(item.get("students") or [])
        enriched.append(row)
    return {"items": enriched}


@router.post("/classes")
async def create_class(
    req: ClassCreateRequest,
    user: dict = Depends(require_roles("admin")),
):
    if not req.name.strip():
        raise HTTPException(status_code=400, detail="班级名称不能为空")
    class_types = await load_class_types()
    if _find_by_id(class_types, req.class_type_id) is None:
        raise HTTPException(status_code=404, detail="班型不存在")
    item = {
        "id": new_entity_id("class"),
        "name": req.name.strip(),
        "class_type_id": req.class_type_id,
        "description": req.description.strip(),
        "teacher_usernames": _dedupe(req.teacher_usernames),
        "students": parse_roster_text(req.roster_text),
        "created_at": now_iso(),
        "created_by": user["username"],
    }
    items = await load_classes()
    items.append(item)
    await save_classes(items)
    return item


@router.post("/users")
async def create_managed_user(
    req: ManagedUserCreateRequest,
    user: dict = Depends(require_roles("admin")),
):
    username = req.username.strip()
    real_name = req.real_name.strip()
    if not username:
        raise HTTPException(status_code=400, detail="用户名不能为空")
    if not real_name:
        raise HTTPException(status_code=400, detail="真实姓名不能为空")
    if await read_json(get_user_path(username)) is not None:
        raise HTTPException(status_code=409, detail="用户名已存在")
    item = _default_user(username, role=req.role, real_name=real_name)
    if req.password:
        item["password_hash"] = _hash_password(req.password)
    item["class_ids"] = _dedupe(req.class_ids)
    item["class_type_ids"] = _dedupe(req.class_type_ids)
    item["expires_at"] = req.expires_at
    item["status"] = "active" if req.activate_now else "pending"
    item["approval_status"] = "approved" if req.activate_now else "pending"
    await write_json(get_user_path(username), item)
    await write_json(get_user_config_path(username), _default_config(username))
    return {"message": "用户已创建", "user": item}


@router.get("/users")
async def list_managed_users(user: dict = Depends(require_roles("admin"))):
    items = []
    for filename in os.listdir(USERS_DIR):
        if not filename.endswith(".json") or filename.endswith("_config.json"):
            continue
        data = await read_json(os.path.join(USERS_DIR, filename))
        if not isinstance(data, dict) or not data.get("username"):
            continue
        if filename != f"{data.get('username')}.json":
            continue
        items.append({
            "username": data.get("username"),
            "real_name": data.get("real_name") or data.get("username"),
            "role": data.get("role", "teacher"),
            "status": data.get("status", "active"),
            "approval_status": data.get("approval_status", "approved"),
            "class_ids": data.get("class_ids") or [],
            "expires_at": data.get("expires_at"),
        })
    items.sort(key=lambda item: (item["role"], item["username"]))
    return {"items": items}


@router.get("/registrations")
async def list_pending_registrations(user: dict = Depends(require_roles("admin"))):
    items = []
    for filename in os.listdir(USERS_DIR):
        if not filename.endswith(".json") or filename.endswith("_config.json"):
            continue
        data = await read_json(os.path.join(USERS_DIR, filename))
        if not isinstance(data, dict) or data.get("role") != "student":
            continue
        if filename != f"{data.get('username')}.json":
            continue
        registration = data.get("registration") or {}
        if data.get("approval_status") == "approved" and registration.get("face_video_path"):
            registration = _delete_registration_video(registration, "approved_legacy_cleanup")
            data["registration"] = registration
            await write_json(os.path.join(USERS_DIR, filename), data)
        items.append({
            "username": data.get("username"),
            "real_name": data.get("real_name") or data.get("username"),
            "status": data.get("status", "pending"),
            "approval_status": data.get("approval_status", "pending"),
            "class_ids": data.get("class_ids") or [],
            "expires_at": data.get("expires_at"),
            "submitted_at": registration.get("submitted_at"),
            "captcha_verified": registration.get("captcha_verified", False),
            "human_check_value": registration.get("human_check_value", ""),
            "face_video_name": registration.get("face_video_name", ""),
            "face_video_available": bool(registration.get("face_video_path") and os.path.exists(registration.get("face_video_path", ""))),
            "face_video_deleted_at": registration.get("face_video_deleted_at"),
            "face_video_delete_reason": registration.get("face_video_delete_reason", ""),
            "face_detection_supported": registration.get("face_detection_supported", False),
            "face_aligned": registration.get("face_aligned", False),
            "face_consent_accepted": registration.get("face_consent_accepted", False),
            "face_consent_version": registration.get("face_consent_version", ""),
            "face_consent_at": registration.get("face_consent_at"),
            "review_notes": registration.get("review_notes", ""),
            "reviewed_at": registration.get("reviewed_at"),
        })
    items.sort(key=lambda item: item.get("submitted_at") or "", reverse=True)
    return {"items": items}


@router.post("/registrations/{username}/review")
async def review_registration(
    username: str,
    req: RegistrationReviewRequest,
    user: dict = Depends(require_roles("admin")),
):
    user_path = get_user_path(username)
    target = await read_json(user_path)
    if not isinstance(target, dict) or target.get("role") != "student":
        raise HTTPException(status_code=404, detail="注册申请不存在")
    target["class_ids"] = _dedupe(req.class_ids) or (target.get("class_ids") or [])
    target["expires_at"] = req.expires_at or target.get("expires_at")
    registration = target.get("registration") or {}
    registration["review_notes"] = req.note.strip()
    registration["reviewed_at"] = now_iso()
    if req.action == "approve":
        target["status"] = "active"
        target["approval_status"] = "approved"
        registration = _delete_registration_video(registration, "approved")
        classes = await load_classes()
        classes, synced_count = sync_student_to_classes(classes, target, target.get("class_ids") or [])
        if synced_count:
            await save_classes(classes)
    else:
        target["status"] = "rejected"
        target["approval_status"] = "rejected"
    target["registration"] = registration
    await write_json(user_path, target)
    return {"message": "审核已更新", "user": target}


@router.get("/distributions")
async def list_distributions(user: dict = Depends(get_current_user_profile)):
    classes = await load_classes()
    class_types = await load_class_types()
    class_map = {item["id"]: item for item in classes}
    class_type_map = {item["id"]: item for item in class_types}
    items = await load_distributions()
    items = _filter_visible_distributions(items, user)
    enriched = []
    for item in items:
        row = dict(item)
        row["stats"] = _compute_distribution_stats(item)
        row["target_names"] = []
        for target_id in item.get("target_ids") or []:
            if item.get("target_type") == "class":
                row["target_names"].append(class_map.get(target_id, {}).get("name", target_id))
            else:
                row["target_names"].append(class_type_map.get(target_id, {}).get("name", target_id))
        enriched.append(row)
    enriched.sort(key=lambda item: item.get("assigned_at", ""), reverse=True)
    return {"items": enriched}


@router.post("/distributions")
async def create_distribution(
    req: DistributionCreateRequest,
    user: dict = Depends(require_roles("admin", "teacher")),
):
    material_ids = _dedupe(req.material_ids)
    target_ids = _dedupe(req.target_ids)
    if not material_ids:
        raise HTTPException(status_code=400, detail="请至少选择一份资料")
    if not target_ids:
        raise HTTPException(status_code=400, detail="请至少选择一个目标")

    class_types = await load_class_types()
    classes = await load_classes()
    students = []
    if req.target_type == "class":
        target_classes = [item for item in classes if item.get("id") in target_ids]
        if user.get("role") != "admin":
            target_classes = [item for item in target_classes if _class_visible_to_user(item, user)]
        if len(target_classes) != len(target_ids):
            raise HTTPException(status_code=403, detail="存在无权分发的班级")
        for item in target_classes:
            students.extend(item.get("students") or [])
    else:
        valid_types = [item for item in class_types if item.get("id") in target_ids]
        if len(valid_types) != len(target_ids):
            raise HTTPException(status_code=404, detail="班型不存在")
        for item in classes:
            if item.get("class_type_id") in target_ids:
                if user.get("role") == "admin" or _class_visible_to_user(item, user):
                    students.extend(item.get("students") or [])

    item = {
        "id": new_entity_id("dist"),
        "material_ids": material_ids,
        "target_type": req.target_type,
        "target_ids": target_ids,
        "tag_path": req.tag_path.strip(),
        "note": req.note.strip(),
        "students": normalize_students(students),
        "completion_records": [],
        "wrong_book_records": [],
        "auto_grade_records": [],
        "explanations": {},
        "assigned_by": user["username"],
        "assigned_at": now_iso(),
    }
    items = await load_distributions()
    items.append(item)
    await save_distributions(items)
    return item


@router.post("/distributions/{distribution_id}/completion")
async def mark_completion(
    distribution_id: str,
    req: CompletionMarkRequest,
    user: dict = Depends(require_roles("admin", "teacher")),
):
    items = await load_distributions()
    target = _find_by_id(items, distribution_id)
    if target is None:
        raise HTTPException(status_code=404, detail="分发记录不存在")
    if not _distribution_visible_to_user(target, user):
        raise HTTPException(status_code=403, detail="无权修改该分发记录")
    target.setdefault("completion_records", []).append({
        "student_name": req.student_name.strip(),
        "completed_parts": [item.strip() for item in req.completed_parts if item.strip()],
        "note": req.note.strip(),
        "marked_by": user["username"],
        "marked_at": now_iso(),
    })
    await save_distributions(items)
    return {"message": "完成情况已记录", "stats": _compute_distribution_stats(target)}


@router.post("/distributions/{distribution_id}/wrong-book")
async def mark_wrong_book(
    distribution_id: str,
    req: WrongBookMarkRequest,
    user: dict = Depends(require_roles("admin", "teacher")),
):
    items = await load_distributions()
    target = _find_by_id(items, distribution_id)
    if target is None:
        raise HTTPException(status_code=404, detail="分发记录不存在")
    if not _distribution_visible_to_user(target, user):
        raise HTTPException(status_code=403, detail="无权修改该分发记录")
    target.setdefault("wrong_book_records", []).append({
        "student_name": req.student_name.strip(),
        "problem_ref": req.problem_ref.strip(),
        "is_wrong": req.is_wrong,
        "note": req.note.strip(),
        "marked_by": user["username"],
        "marked_at": now_iso(),
    })
    await save_distributions(items)
    return {"message": "错题本已更新", "stats": _compute_distribution_stats(target)}


@router.post("/distributions/{distribution_id}/auto-grade")
async def auto_grade_distribution(
    distribution_id: str,
    req: AutoGradeRequest,
    user: dict = Depends(require_roles("admin", "teacher")),
):
    student_name = req.student_name.strip()
    answer_material_id = req.answer_material_id.strip()
    submission_material_id = req.submission_material_id.strip()
    if not student_name:
        raise HTTPException(status_code=400, detail="请选择学生")
    if not answer_material_id or not submission_material_id:
        raise HTTPException(status_code=400, detail="请同时选择标准答案和学生作答")
    if answer_material_id == submission_material_id:
        raise HTTPException(status_code=400, detail="标准答案和学生作答不能是同一份资料")

    items = await load_distributions()
    target = _find_by_id(items, distribution_id)
    if target is None:
        raise HTTPException(status_code=404, detail="分发记录不存在")
    if not _distribution_visible_to_user(target, user):
        raise HTTPException(status_code=403, detail="无权修改该分发记录")
    student_names = {
        str(row.get("name") or row.get("real_name") or "").strip()
        for row in (target.get("students") or [])
        if isinstance(row, dict)
    }
    if student_name not in student_names:
        raise HTTPException(status_code=400, detail="该学生不在当前分发名单中")

    subject = await get_user_subject(user["username"])
    materials = await read_json(get_materials_path(user["username"], subject)) or []
    material_map = {str(row.get("id")): row for row in materials if isinstance(row, dict) and row.get("id")}
    answer_material = material_map.get(answer_material_id)
    submission_material = material_map.get(submission_material_id)
    if answer_material is None:
        raise HTTPException(status_code=404, detail="找不到标准答案资料")
    if submission_material is None:
        raise HTTPException(status_code=404, detail="找不到学生作答资料")

    result = await _run_kimi_auto_grade(answer_material, submission_material, user)
    grade_record, created_count = _apply_auto_grade_result(target, req, result, user)
    await save_distributions(items)
    return {
        "message": f"AI 已完成批改，自动登记 {created_count} 条错题",
        "grade": grade_record,
        "stats": _compute_distribution_stats(target),
    }


@router.post("/distributions/{distribution_id}/explanations")
async def generate_wrong_book_explanation(
    distribution_id: str,
    req: ExplanationCreateRequest,
    user: dict = Depends(require_roles("admin", "teacher")),
):
    problem_ref = req.problem_ref.strip()
    if not problem_ref:
        raise HTTPException(status_code=400, detail="题目标识不能为空")
    items = await load_distributions()
    target = _find_by_id(items, distribution_id)
    if target is None:
        raise HTTPException(status_code=404, detail="分发记录不存在")
    if not _distribution_visible_to_user(target, user):
        raise HTTPException(status_code=403, detail="无权修改该分发记录")
    has_wrong_record = any(
        row.get("is_wrong") and (row.get("problem_ref") or "") == problem_ref
        for row in (target.get("wrong_book_records") or [])
    )
    if not has_wrong_record:
        raise HTTPException(status_code=404, detail="该题尚未登记错题")
    explanations = target.setdefault("explanations", {})
    if not req.force and isinstance(explanations.get(problem_ref), dict):
        return {"message": "解析已存在", "explanation": explanations[problem_ref], "stats": _compute_distribution_stats(target)}

    context = _wrong_book_context(target, problem_ref)
    explanation = await _generate_kimi_explanation(problem_ref, context, user)
    explanations[problem_ref] = explanation
    target["explanations"] = explanations
    await save_distributions(items)
    return {"message": "AI 解析已生成", "explanation": explanation, "stats": _compute_distribution_stats(target)}


@router.get("/distributions/{distribution_id}/exports/wrong-book.csv")
async def export_distribution_wrong_book(
    distribution_id: str,
    student_name: str = "",
    user: dict = Depends(require_roles("admin", "teacher")),
):
    items = _filter_visible_distributions(await load_distributions(), user)
    target = _find_by_id(items, distribution_id)
    if target is None:
        raise HTTPException(status_code=404, detail="分发记录不存在")

    filtered_rows = []
    explanations = target.get("explanations") or {}
    for row in target.get("wrong_book_records") or []:
        if not row.get("is_wrong"):
            continue
        if student_name and row.get("student_name") != student_name:
            continue
        explanation = explanations.get(row.get("problem_ref", "")) if isinstance(explanations, dict) else None
        filtered_rows.append([
            row.get("student_name", ""),
            row.get("problem_ref", ""),
            row.get("note", ""),
            explanation.get("content", "") if isinstance(explanation, dict) else "",
            row.get("marked_at", ""),
            target.get("tag_path", ""),
            target.get("note", ""),
        ])

    filename = f"wrong-book-{distribution_id}"
    if student_name:
        filename += f"-{student_name}"
    filename += ".csv"
    return _csv_response(
        filename,
        ["学生", "题目标识", "备注", "AI解析", "登记时间", "标签路径", "分发备注"],
        filtered_rows,
    )


@router.get("/distributions/{distribution_id}/exports/wrong-book.docx.zip")
async def export_distribution_wrong_book_docx(
    distribution_id: str,
    student_name: str = "",
    user: dict = Depends(require_roles("admin", "teacher")),
):
    items = _filter_visible_distributions(await load_distributions(), user)
    target = _find_by_id(items, distribution_id)
    if target is None:
        raise HTTPException(status_code=404, detail="分发记录不存在")
    entries = _collect_wrong_book_entries([target], student_name)
    title = "错题本"
    if student_name:
        title += f"_{student_name}"
    else:
        title += f"_{distribution_id}"
    return _wrong_book_docx_zip_response(entries, title, f"{title}.zip")


@router.get("/exports/wrong-book.csv")
async def export_all_wrong_books(
    student_name: str = "",
    user: dict = Depends(require_roles("admin", "teacher")),
):
    items = _filter_visible_distributions(await load_distributions(), user)
    rows = []
    for item in items:
        explanations = item.get("explanations") or {}
        for row in item.get("wrong_book_records") or []:
            if not row.get("is_wrong"):
                continue
            if student_name and row.get("student_name") != student_name:
                continue
            explanation = explanations.get(row.get("problem_ref", "")) if isinstance(explanations, dict) else None
            rows.append([
                row.get("student_name", ""),
                row.get("problem_ref", ""),
                row.get("note", ""),
                explanation.get("content", "") if isinstance(explanation, dict) else "",
                row.get("marked_at", ""),
                item.get("tag_path", ""),
                item.get("note", ""),
                item.get("assigned_at", ""),
            ])
    filename = "wrong-book-all.csv" if not student_name else f"wrong-book-{student_name}.csv"
    return _csv_response(
        filename,
        ["学生", "题目标识", "备注", "AI解析", "登记时间", "标签路径", "分发备注", "分发时间"],
        rows,
    )


@router.get("/exports/wrong-book.docx.zip")
async def export_all_wrong_books_docx(
    student_name: str = "",
    user: dict = Depends(require_roles("admin", "teacher")),
):
    items = _filter_visible_distributions(await load_distributions(), user)
    entries = _collect_wrong_book_entries(items, student_name)
    title = "整合错题本"
    if student_name:
        title += f"_{student_name}"
    return _wrong_book_docx_zip_response(entries, title, f"{title}.zip")


@router.get("/distributions/{distribution_id}/exports/correctness-summary.csv")
async def export_distribution_correctness_summary(
    distribution_id: str,
    user: dict = Depends(require_roles("admin", "teacher")),
):
    items = _filter_visible_distributions(await load_distributions(), user)
    target = _find_by_id(items, distribution_id)
    if target is None:
        raise HTTPException(status_code=404, detail="分发记录不存在")
    stats = _compute_distribution_stats(target)
    rows = []
    for row in stats["problem_stats"]:
        rows.append([
            row["problem_ref"],
            stats["student_total"],
            row["wrong_count"],
            f"{row['correct_rate'] * 100:.2f}%",
            "、".join(row["wrong_students"]),
            row.get("explanation", ""),
            target.get("tag_path", ""),
        ])
    return _csv_response(
        f"correctness-summary-{distribution_id}.csv",
        ["题目标识", "学生总数", "错误人数", "正确率", "错误学生", "AI解析", "标签路径"],
        rows,
    )


@router.get("/exports/correctness-summary.csv")
async def export_all_correctness_summary(
    user: dict = Depends(require_roles("admin", "teacher")),
):
    items = _filter_visible_distributions(await load_distributions(), user)
    rows = []
    for item in items:
        stats = _compute_distribution_stats(item)
        for row in stats["problem_stats"]:
            rows.append([
                item.get("id", ""),
                row["problem_ref"],
                stats["student_total"],
                row["wrong_count"],
                f"{row['correct_rate'] * 100:.2f}%",
                "、".join(row["wrong_students"]),
                row.get("explanation", ""),
                item.get("tag_path", ""),
                item.get("assigned_at", ""),
            ])
    return _csv_response(
        "correctness-summary-all.csv",
        ["分发ID", "题目标识", "学生总数", "错误人数", "正确率", "错误学生", "AI解析", "标签路径", "分发时间"],
        rows,
    )
