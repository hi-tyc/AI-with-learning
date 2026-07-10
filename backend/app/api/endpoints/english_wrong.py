from fastapi import APIRouter, HTTPException, Depends, Query
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional
from datetime import datetime, timedelta
import uuid
import io

from app.api.endpoints.auth import get_current_user
from app.core.user_data import (
    get_user_subject,
    get_problems_path,
    get_answers_path,
    get_wrong_english_path,
    get_problem_answer_map_path,
    get_words_path,
)
from app.utils.file_lock import read_json, write_json

router = APIRouter()


class AddWrongItem(BaseModel):
    problem_ids: list[str]
    source_file: str = ""
    tag: str = ""
    exam_tag: str = ""


class RemoveWrongRequest(BaseModel):
    ids: list[str]


class AddWrongWordsRequest(BaseModel):
    word_list_id: str
    word_indices: list[int]
    word_list_name: str = ""
    tag: str = ""


@router.post("/add")
async def add_wrong_items(
    req: AddWrongItem,
    username: str = Depends(get_current_user),
):
    subject = await get_user_subject(username)
    problems = await read_json(get_problems_path(username, subject)) or []
    answers = await read_json(get_answers_path(username, subject)) or []
    mapping = await read_json(get_problem_answer_map_path(username, subject)) or {}
    p2a = mapping.get("problem_id_to_answer_ids", {})

    wrong_list = await read_json(get_wrong_english_path(username, subject)) or []
    existing_ids = {w["problem_id"] for w in wrong_list if w.get("type", "problem") == "problem"}

    now = datetime.now().isoformat()
    added = 0
    for pid in req.problem_ids:
        if pid in existing_ids:
            continue
        problem = next((p for p in problems if p["id"] == pid), None)
        if not problem:
            continue
        content = problem.get("content", "")
        source_file = req.source_file or problem.get("source_file", "") or problem.get("filename", "")
        tag = req.tag or problem.get("exam", "")
        exam_tag = req.exam_tag or tag

        answer_ids = p2a.get(pid, [])
        answer_contents = []
        mismatch_flag = not bool(answer_ids)
        for aid in answer_ids:
            ans = next((a for a in answers if a["id"] == aid), None)
            if ans:
                answer_contents.append(ans.get("content", ""))

        wrong_list.append({
            "id": str(uuid.uuid4())[:8],
            "type": "problem",
            "problem_id": pid,
            "content": content,
            "answer_ids": answer_ids,
            "answer_contents": answer_contents,
            "mismatch_flag": mismatch_flag,
            "source_file": source_file,
            "tag": tag,
            "exam_tag": exam_tag,
            "created_at": now,
        })
        existing_ids.add(pid)
        added += 1

    await write_json(get_wrong_english_path(username, subject), wrong_list)
    return {"added": added, "total": len(wrong_list)}


@router.post("/add-words")
async def add_wrong_words(
    req: AddWrongWordsRequest,
    username: str = Depends(get_current_user),
):
    subject = await get_user_subject(username)
    words_list = await read_json(get_words_path(username, subject)) or []
    word_list = next((w for w in words_list if w["id"] == req.word_list_id), None)
    if not word_list:
        raise HTTPException(status_code=404, detail="词单不存在")

    all_words = word_list.get("words", [])
    if not all_words and word_list.get("is_student"):
        # Try teacher version
        teacher_id = word_list.get("matched_teacher_id")
        if teacher_id:
            tw = next((w for w in words_list if w["id"] == teacher_id), None)
            if tw:
                all_words = tw.get("words", [])

    wrong_list = await read_json(get_wrong_english_path(username, subject)) or []
    existing_keys = {(w.get("word_list_id", ""), w.get("word_index", -1)) for w in wrong_list if w.get("type") == "word"}

    now = datetime.now().isoformat()
    added = 0
    for idx in req.word_indices:
        if idx < 0 or idx >= len(all_words):
            continue
        key = (req.word_list_id, idx)
        if key in existing_keys:
            continue
        w = all_words[idx]
        wrong_list.append({
            "id": str(uuid.uuid4())[:8],
            "type": "word",
            "word_list_id": req.word_list_id,
            "word_index": idx,
            "word_english": w.get("english", ""),
            "word_chinese": w.get("chinese", ""),
            "word_pos": w.get("pos", ""),
            "word_list_name": req.word_list_name or word_list.get("filename", ""),
            "source_file": word_list.get("filename", ""),
            "tag": req.tag or word_list.get("tag", ""),
            "exam_tag": word_list.get("tag", ""),
            "created_at": now,
        })
        existing_keys.add(key)
        added += 1

    await write_json(get_wrong_english_path(username, subject), wrong_list)
    return {"added": added, "total": len(wrong_list)}


@router.get("/list")
async def list_wrong_items(
    username: str = Depends(get_current_user),
    time_range: Optional[str] = Query(None),
    exam_tag: Optional[str] = Query(None),
    show_all: bool = Query(True),
    item_type: Optional[str] = Query(None),
):
    subject = await get_user_subject(username)
    wrong_list = await read_json(get_wrong_english_path(username, subject)) or []

    if item_type:
        wrong_list = [w for w in wrong_list if w.get("type", "problem") == item_type]

    if not show_all:
        wrong_list = [w for w in wrong_list if w.get("mismatch_flag")]

    if exam_tag:
        wrong_list = [w for w in wrong_list if w.get("exam_tag") == exam_tag]

    if time_range:
        now = datetime.now()
        if time_range == "today":
            start = now.replace(hour=0, minute=0, second=0, microsecond=0)
        elif time_range == "week":
            start = now - timedelta(days=now.weekday())
            start = start.replace(hour=0, minute=0, second=0, microsecond=0)
        else:
            start = None
        if start:
            wrong_list = [w for w in wrong_list if datetime.fromisoformat(w.get("created_at", "")) >= start]

    problems = await read_json(get_problems_path(username, subject)) or []
    prob_map = {p["id"]: p for p in problems}
    for w in wrong_list:
        if w.get("type", "problem") == "problem":
            p = prob_map.get(w["problem_id"])
            w["has_solution"] = bool(p and p.get("solution"))
        else:
            w["has_solution"] = False

    wrong_list.reverse()
    return {"items": wrong_list, "total": len(wrong_list)}


@router.get("/exam-tags")
async def get_exam_tags(
    username: str = Depends(get_current_user),
):
    subject = await get_user_subject(username)
    wrong_list = await read_json(get_wrong_english_path(username, subject)) or []
    tags = sorted(set(w.get("exam_tag", "") for w in wrong_list if w.get("exam_tag")))
    return {"tags": tags}


@router.post("/remove")
async def remove_wrong_items(
    req: RemoveWrongRequest,
    username: str = Depends(get_current_user),
):
    subject = await get_user_subject(username)
    wrong_list = await read_json(get_wrong_english_path(username, subject)) or []
    remove_set = set(req.ids)
    new_list = [w for w in wrong_list if w["id"] not in remove_set]
    await write_json(get_wrong_english_path(username, subject), new_list)
    return {"removed": len(req.ids), "remaining": len(new_list)}


@router.post("/toggle-problem/{problem_id}")
async def toggle_problem_wrong(
    problem_id: str,
    username: str = Depends(get_current_user),
):
    subject = await get_user_subject(username)
    problems = await read_json(get_problems_path(username, subject)) or []
    answers = await read_json(get_answers_path(username, subject)) or []
    mapping = await read_json(get_problem_answer_map_path(username, subject)) or {}
    p2a = mapping.get("problem_id_to_answer_ids", {})

    wrong_list = await read_json(get_wrong_english_path(username, subject)) or []
    existing = [w for w in wrong_list if w.get("problem_id") == problem_id and w.get("type", "problem") == "problem"]

    if existing:
        new_list = [w for w in wrong_list if w["id"] not in {e["id"] for e in existing}]
        await write_json(get_wrong_english_path(username, subject), new_list)
        return {"message": "已移出错题本", "is_wrong": False}

    problem = next((p for p in problems if p["id"] == problem_id), None)
    if not problem:
        raise HTTPException(status_code=404, detail="题目不存在")

    answer_ids = p2a.get(problem_id, [])
    answer_contents = []
    mismatch_flag = not bool(answer_ids)
    for aid in answer_ids:
        ans = next((a for a in answers if a["id"] == aid), None)
        if ans:
            answer_contents.append(ans.get("content", ""))

    wrong_list.append({
        "id": str(uuid.uuid4())[:8],
        "type": "problem",
        "problem_id": problem_id,
        "content": problem.get("content", ""),
        "answer_ids": answer_ids,
        "answer_contents": answer_contents,
        "mismatch_flag": mismatch_flag,
        "source_file": problem.get("source_file", "") or problem.get("filename", ""),
        "tag": problem.get("exam", ""),
        "exam_tag": problem.get("exam", ""),
        "created_at": datetime.now().isoformat(),
    })
    await write_json(get_wrong_english_path(username, subject), wrong_list)
    return {"message": "已加入错题本", "is_wrong": True, "mismatch_flag": mismatch_flag}


@router.post("/toggle-answer/{answer_id}")
async def toggle_answer_wrong(
    answer_id: str,
    username: str = Depends(get_current_user),
):
    subject = await get_user_subject(username)
    answers = await read_json(get_answers_path(username, subject)) or []
    mapping = await read_json(get_problem_answer_map_path(username, subject)) or {}
    p2a = mapping.get("problem_id_to_answer_ids", {})
    a2p = mapping.get("answer_id_to_problem_ids", {})

    answer = next((a for a in answers if a["id"] == answer_id), None)
    if not answer:
        raise HTTPException(status_code=404, detail="答案不存在")

    problem_ids = a2p.get(answer_id, [])

    if problem_ids:
        pid = problem_ids[0]
        problems = await read_json(get_problems_path(username, subject)) or []
        wrong_list = await read_json(get_wrong_english_path(username, subject)) or []
        existing = [w for w in wrong_list if w.get("problem_id") == pid and w.get("type", "problem") == "problem"]

        if existing:
            new_list = [w for w in wrong_list if w["id"] not in {e["id"] for e in existing}]
            await write_json(get_wrong_english_path(username, subject), new_list)
            return {"message": "已移出错题本", "is_wrong": False}

        problem = next((p for p in problems if p["id"] == pid), None)
        if problem:
            answer_ids = p2a.get(pid, [])
            answer_contents = []
            for aid in answer_ids:
                ans = next((a for a in answers if a["id"] == aid), None)
                if ans:
                    answer_contents.append(ans.get("content", ""))

            wrong_list.append({
                "id": str(uuid.uuid4())[:8],
                "type": "problem",
                "problem_id": pid,
                "content": problem.get("content", ""),
                "answer_ids": answer_ids,
                "answer_contents": answer_contents,
                "mismatch_flag": False,
                "source_file": problem.get("source_file", "") or problem.get("filename", ""),
                "tag": problem.get("exam", ""),
                "exam_tag": problem.get("exam", ""),
                "created_at": datetime.now().isoformat(),
            })
            await write_json(get_wrong_english_path(username, subject), wrong_list)
            return {"message": "已加入错题本", "is_wrong": True}
    else:
        wrong_list = await read_json(get_wrong_english_path(username, subject)) or []
        existing = [w for w in wrong_list if w.get("type") == "answer" and w.get("answer_id") == answer_id]

        if existing:
            new_list = [w for w in wrong_list if w["id"] not in {e["id"] for e in existing}]
            await write_json(get_wrong_english_path(username, subject), new_list)
            return {"message": "已移出错题本", "is_wrong": False}

        wrong_list.append({
            "id": str(uuid.uuid4())[:8],
            "type": "answer",
            "answer_id": answer_id,
            "content": answer.get("content", ""),
            "answer_contents": [],
            "mismatch_flag": True,
            "source_file": answer.get("source_file", "") or answer.get("filename", ""),
            "tag": answer.get("tag", ""),
            "exam_tag": answer.get("tag", ""),
            "created_at": datetime.now().isoformat(),
        })
        await write_json(get_wrong_english_path(username, subject), wrong_list)
        return {"message": "已加入错题本（未匹配题目）", "is_wrong": True, "mismatch_flag": True}

    raise HTTPException(status_code=404, detail="操作失败")


@router.post("/toggle-word/{word_list_id}/{word_index}")
async def toggle_word_wrong(
    word_list_id: str,
    word_index: int,
    username: str = Depends(get_current_user),
):
    subject = await get_user_subject(username)
    words_list = await read_json(get_words_path(username, subject)) or []
    word_list = next((w for w in words_list if w["id"] == word_list_id), None)
    if not word_list:
        raise HTTPException(status_code=404, detail="词单不存在")

    wrong_list = await read_json(get_wrong_english_path(username, subject)) or []
    existing = [w for w in wrong_list if w.get("type") == "word" and w.get("word_list_id") == word_list_id and w.get("word_index") == word_index]

    if existing:
        new_list = [w for w in wrong_list if w["id"] not in {e["id"] for e in existing}]
        await write_json(get_wrong_english_path(username, subject), new_list)
        return {"message": "已移出错题本", "is_wrong": False}

    all_words = word_list.get("words", [])
    if not all_words and word_list.get("is_student"):
        teacher_id = word_list.get("matched_teacher_id")
        if teacher_id:
            tw = next((w for w in words_list if w["id"] == teacher_id), None)
            if tw:
                all_words = tw.get("words", [])

    if word_index < 0 or word_index >= len(all_words):
        raise HTTPException(status_code=400, detail="单词索引超出范围")

    w = all_words[word_index]
    wrong_list.append({
        "id": str(uuid.uuid4())[:8],
        "type": "word",
        "word_list_id": word_list_id,
        "word_index": word_index,
        "word_english": w.get("english", ""),
        "word_chinese": w.get("chinese", ""),
        "word_pos": w.get("pos", ""),
        "word_list_name": word_list.get("filename", ""),
        "source_file": word_list.get("filename", ""),
        "tag": word_list.get("tag", ""),
        "exam_tag": word_list.get("tag", ""),
        "created_at": datetime.now().isoformat(),
    })
    await write_json(get_wrong_english_path(username, subject), wrong_list)
    return {"message": "已加入错题本", "is_wrong": True}


@router.get("/check/{problem_id}")
async def check_wrong_status(
    problem_id: str,
    username: str = Depends(get_current_user),
):
    subject = await get_user_subject(username)
    wrong_list = await read_json(get_wrong_english_path(username, subject)) or []
    is_wrong = any(w["problem_id"] == problem_id for w in wrong_list if w.get("type", "problem") == "problem")
    return {"is_wrong": is_wrong}


@router.get("/check-answer/{answer_id}")
async def check_answer_wrong(
    answer_id: str,
    username: str = Depends(get_current_user),
):
    subject = await get_user_subject(username)
    mapping = await read_json(get_problem_answer_map_path(username, subject)) or {}
    a2p = mapping.get("answer_id_to_problem_ids", {})
    problem_ids = a2p.get(answer_id, [])

    if problem_ids:
        pid = problem_ids[0]
        wrong_list = await read_json(get_wrong_english_path(username, subject)) or []
        is_wrong = any(w.get("problem_id") == pid and w.get("type", "problem") == "problem" for w in wrong_list)
        return {"is_wrong": is_wrong}
    else:
        wrong_list = await read_json(get_wrong_english_path(username, subject)) or []
        is_wrong = any(w.get("type") == "answer" and w.get("answer_id") == answer_id for w in wrong_list)
        return {"is_wrong": is_wrong}


@router.get("/check-word/{word_list_id}/{word_index}")
async def check_word_wrong(
    word_list_id: str,
    word_index: int,
    username: str = Depends(get_current_user),
):
    subject = await get_user_subject(username)
    wrong_list = await read_json(get_wrong_english_path(username, subject)) or []
    is_wrong = any(w.get("type") == "word" and w.get("word_list_id") == word_list_id and w.get("word_index") == word_index for w in wrong_list)
    return {"is_wrong": is_wrong}


@router.post("/generate")
async def generate_wrong_book(
    data: dict,
    username: str = Depends(get_current_user),
):
    time_range = data.get("time_range", "all")
    exam_tag = data.get("exam_tag", "")
    show_all = data.get("show_all", True)
    gen_type = data.get("gen_type", "all")

    subject = await get_user_subject(username)
    wrong_list = await read_json(get_wrong_english_path(username, subject)) or []

    if not show_all:
        wrong_list = [w for w in wrong_list if w.get("mismatch_flag")]

    if exam_tag:
        wrong_list = [w for w in wrong_list if w.get("exam_tag") == exam_tag]

    if time_range and time_range != "all":
        now = datetime.now()
        if time_range == "today":
            start = now.replace(hour=0, minute=0, second=0, microsecond=0)
        elif time_range == "week":
            start = now - timedelta(days=now.weekday())
            start = start.replace(hour=0, minute=0, second=0, microsecond=0)
        else:
            start = None
        if start:
            wrong_list = [w for w in wrong_list if datetime.fromisoformat(w.get("created_at", "")) >= start]

    if gen_type == "word":
        wrong_list = [w for w in wrong_list if w.get("type") == "word"]
    elif gen_type == "problem":
        wrong_list = [w for w in wrong_list if w.get("type") in ("problem", "answer")]

    # 词单排最前
    wrong_list.sort(key=lambda x: (0 if x.get("type") == "word" else 1, 0))

    problems = await read_json(get_problems_path(username, subject)) or []
    prob_map = {p["id"]: p for p in problems}

    async def generate():
        try:
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        except ImportError:
            yield {"type": "error", "text": "生成Word文档失败: 缺少 python-docx 库，请执行 pip install python-docx"}
            return
        try:

            doc_student = Document()
            doc_answer = Document()

            for doc in (doc_student, doc_answer):
                section = doc.sections[0]
                section.top_margin = Cm(2.5)
                section.bottom_margin = Cm(2.5)
                section.left_margin = Cm(2.5)
                section.right_margin = Cm(2.5)
                pfmt = doc.styles['Normal'].paragraph_format
                pfmt.space_before = Pt(0)
                pfmt.space_after = Pt(0)
                pfmt.line_spacing = Pt(20)
                pfmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                doc.styles['Normal'].font.size = Pt(10.5)

            title_text = "英语错题本"
            if exam_tag:
                title_text += f" - {exam_tag}"
            if time_range == "today":
                title_text += "（今日标记）"
            elif time_range == "week":
                title_text += "（本周标记）"

            for doc in (doc_student, doc_answer):
                heading = doc.add_heading(title_text, level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                heading.runs[0].font.size = Pt(16)
                heading.paragraph_format.space_after = Pt(8)

                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"共 {len(wrong_list)} 项  |  生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                p.paragraph_format.space_after = Pt(8)

            def _add_choice_row(d, parts):
                """Format choice options in a single table row with 4 evenly-spaced columns."""
                tbl = d.add_table(rows=1, cols=len(parts))
                tbl.autofit = False
                tbl.style = 'Table Grid'
                col_width = Cm(14.4 / max(len(parts), 4))
                for c in tbl.columns:
                    c.width = col_width
                for ci, part in enumerate(parts):
                    cell = tbl.cell(0, ci)
                    p = cell.paragraphs[0]
                    run = p.add_run(part.strip())
                    run.font.size = Pt(10.5)
                    run.font.name = 'Arial'
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                return tbl

            def _detect_choices(text):
                """Detect multiple choice options (A. ... B. ... etc.) and return list of parts or None."""
                import re
                m = re.match(r'^([A-D])\s*[.、]\s*(.*?)(?=\s+([A-D])\s*[.、]|\s*$)', text, re.DOTALL)
                if not m:
                    return None
                parts = re.split(r'\s+([A-D])\s*[.、]', text)
                result = []
                for i in range(0, len(parts), 2):
                    if i + 1 < len(parts):
                        result.append(f"{parts[i+1]}. {parts[i+2].strip()}" if i+2 < len(parts) else parts[i+1])
                return result if len(result) >= 2 else None

            # Group consecutive words from same word list
            word_groups = []
            current_wg = None
            for w in wrong_list:
                if w.get("type") == "word":
                    wlid = w.get("word_list_id", "")
                    if current_wg is None or current_wg["wlid"] != wlid:
                        current_wg = {"wlid": wlid, "words": []}
                        word_groups.append(current_wg)
                    current_wg["words"].append(w)
                else:
                    current_wg = None

            # Map word index for each word group -> table row
            wg_idx = 0
            word_in_group = 0
            for w in wrong_list:
                if w.get("type") == "word":
                    wg = word_groups[wg_idx]
                    w["_table_row"] = word_in_group
                    word_in_group += 1
                    if word_in_group >= len(wg["words"]):
                        wg_idx += 1
                        word_in_group = 0

            # Render word groups as single tables
            for wg in word_groups:
                for doc, is_answer in [(doc_student, False), (doc_answer, True)]:
                    tbl = doc.add_table(rows=len(wg["words"]), cols=2)
                    tbl.autofit = False
                    tbl.style = 'Table Grid'
                    for c in tbl.columns:
                        c.width = Cm(7.2)
                    for ri, w in enumerate(wg["words"]):
                        eng = w.get("word_english", "")
                        chi = w.get("word_chinese", "")
                        for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                            for ci in range(2):
                                cell = tbl.cell(ri, ci)
                                tc = cell._tc
                                tcPr = tc.get_or_add_tcPr()
                                border = __import__('lxml').etree.SubElement(tcPr, '{%s}%s' % ('http://schemas.openxmlformats.org/wordprocessingml/2006/main', border_name))
                                border.set('{%s}val' % 'http://schemas.openxmlformats.org/wordprocessingml/2006/main', 'nil')
                        is_multi = len(eng.split()) > 1
                        has_blank = "＿" in eng or "_" in eng
                        if is_answer:
                            display_eng = eng
                        else:
                            if has_blank:
                                display_eng = eng
                            elif is_multi:
                                display_eng = "＿" * 20
                            else:
                                display_eng = (eng[:1] if eng else "") + "＿" * 15
                        cell_l = tbl.cell(ri, 0)
                        cell_r = tbl.cell(ri, 1)
                        for cell, txt in [(cell_l, chi), (cell_r, display_eng)]:
                            p = cell.paragraphs[0]
                            run = p.add_run(txt)
                            run.font.size = Pt(10.5)
                            run.font.name = 'Arial'
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after = Pt(0)
                            p.paragraph_format.line_spacing = Pt(20)
                            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                for doc in (doc_student, doc_answer):
                    doc.add_paragraph()

            # Render non-word items
            for idx, w in enumerate(wrong_list):
                if w.get("type") == "word":
                    continue

                q_num = f"{idx + 1}."
                content = w.get("content", "")
                answer_contents = w.get("answer_contents", [])
                mismatch = w.get("mismatch_flag", False)
                pid = w.get("problem_id", "")
                p_solution = prob_map.get(pid, {})
                solution = p_solution.get("solution", "") if p_solution else ""

                for doc, is_answer in [(doc_student, False), (doc_answer, True)]:
                    if not is_answer:
                        p_q = doc.add_paragraph()
                        run_num = p_q.add_run(q_num + " ")
                        run_num.bold = True
                        run_num.font.size = Pt(11)
                        run_num.font.name = 'Arial'
                        choice_parts = _detect_choices(content)
                        if choice_parts:
                            p_q.paragraph_format.space_after = Pt(2)
                            _add_choice_row(doc, choice_parts)
                        else:
                            run_content = p_q.add_run(content)
                            run_content.font.size = Pt(11)
                            run_content.font.name = 'Arial'
                        p_q.paragraph_format.space_before = Pt(0)
                        p_q.paragraph_format.space_after = Pt(0)

                        if mismatch:
                            p_tag = doc.add_paragraph()
                            run_tag = p_tag.add_run("[未匹配]")
                            run_tag.font.size = Pt(9)
                            run_tag.font.color.rgb = RGBColor(0xdc, 0x26, 0x26)
                            run_tag.italic = True
                            p_tag.paragraph_format.space_before = Pt(0)
                            p_tag.paragraph_format.space_after = Pt(0)
                    else:
                        if answer_contents:
                            for ac in answer_contents:
                                p_a = doc.add_paragraph()
                                run_a_val = p_a.add_run(ac)
                                run_a_val.font.size = Pt(11)
                                run_a_val.font.name = 'Arial'
                                p_a.paragraph_format.space_before = Pt(0)
                                p_a.paragraph_format.space_after = Pt(0)
                        elif solution:
                            p_s = doc.add_paragraph()
                            run_s_val = p_s.add_run(solution[:500])
                            run_s_val.font.size = Pt(11)
                            run_s_val.font.name = 'Arial'
                            p_s.paragraph_format.space_before = Pt(0)
                            p_s.paragraph_format.space_after = Pt(0)
                        elif w.get("type") == "answer":
                            p_a = doc.add_paragraph()
                            run_a_val = p_a.add_run(content)
                            run_a_val.font.size = Pt(11)
                            run_a_val.font.name = 'Arial'
                            p_a.paragraph_format.space_before = Pt(0)
                            p_a.paragraph_format.space_after = Pt(0)

                if w.get("type") != "word":
                    for doc in (doc_student, doc_answer):
                        doc.add_paragraph()

            student_buf = io.BytesIO()
            doc_student.save(student_buf)
            student_buf.seek(0)

            answer_buf = io.BytesIO()
            doc_answer.save(answer_buf)
            answer_buf.seek(0)

            yield {
                "type": "result",
                "student_docx": student_buf.getvalue().hex(),
                "answer_docx": answer_buf.getvalue().hex(),
                "count": len(wrong_list),
            }
        except Exception as e:
            yield {"type": "error", "text": f"生成Word文档失败: {str(e)[:200]}"}

    gen = generate()
    async def event_generator():
        async for item in gen:
            if item["type"] == "result":
                yield "data: " + __import__('json').dumps(item, ensure_ascii=False) + "\n\n"
                yield "data: [DONE]\n\n"
            else:
                yield "data: " + __import__('json').dumps(item, ensure_ascii=False) + "\n\n"
                yield "data: [DONE]\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream")
